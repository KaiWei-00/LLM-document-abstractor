"""
Document Information Extraction Microservice using LangGraph
Supports PDF, Excel, and Word files with dynamic schema-based extraction
Enhanced with AI Schema Detection Agent
"""

import os
import re
import json
import robust_json  # Import our robust JSON parsing utilities
import content_trimmer  # Import content trimming utilities
import excel_preprocessing
import excel_fix  # Import the simplified Excel fix module
from excel_fix import ai_excel_document_understanding, ai_excel_schema_based_extraction, describe_schema  # Import AI functions
import shutil
import tempfile
import traceback
import time
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
from datetime import datetime

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# FastAPI imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field

# LangGraph imports
from langgraph.graph import StateGraph, START, END
from langgraph.prebuilt import ToolNode
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langchain_core.tools import tool

# Document processing imports
import PyPDF2
import pandas as pd
import docx
import openpyxl
from io import BytesIO
import fitz  # PyMuPDF for better PDF handling
from pdf_preprocessing import enhance_pdf_extraction  # PDF abstraction utilities
import re  # Regular expressions for pattern matching

# PDF image-based extraction utility
from pdf_image_extraction import extract_pdf_via_images, aggregate_page_results

# Validate environment variables
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("OPENAI_API_KEY environment variable is required. Please set it in your .env file.")

# Initialize FastAPI app
app = FastAPI(title="Document Information Extraction API with Schema Detection", version="1.1.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OpenAI LLM
try:
    # Check if the specified model is a completion model or a chat model
    model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
    api_key = os.getenv("OPENAI_API_KEY")
    
    # Debug logging (with masked key for security)
    masked_key = api_key[:5] + "..." + api_key[-4:] if api_key and len(api_key) > 10 else "None"
    print(f"Debug - Using model: {model_name}, API key: {masked_key}")
    
    if not api_key:
        raise ValueError("API key is missing or empty. Check your .env file.")
        
    if "instruct" in model_name:
        from langchain_openai import OpenAI
        llm = OpenAI(
            model=model_name,
            temperature=0,
            openai_api_key=api_key
        )
        print(f"Successfully initialized OpenAI completion model: {model_name}")
    else:
        # Default to ChatOpenAI for other models
        llm = ChatOpenAI(
            model=model_name,
            temperature=0,
            openai_api_key=api_key
        )
        print(f"Successfully initialized ChatOpenAI model: {model_name}")

    # Test if llm is properly initialized
    print(f"LLM initialized as type: {type(llm).__name__}")
    
except Exception as e:
    print(f"Error initializing OpenAI: {str(e)}")
    import traceback
    print(f"Error traceback: {traceback.format_exc()}")
    llm = None  # Ensure llm is always defined, even on failure

# --- Utility: Flatten/Describe Nested Schema for Prompt Generation ---
def describe_schema(schema, indent=0):
    """
    Recursively describe a (possibly nested) schema as a readable string for LLM prompts.
    Supports dicts, lists, and flat fields.
    """
    lines = []
    prefix = '  ' * indent
    if isinstance(schema, dict):
        for key, value in schema.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}- {key}:")
                lines.extend(describe_schema(value, indent+1))
            elif isinstance(value, list):
                lines.append(f"{prefix}- {key}: (list)")
                if value:
                    lines.extend(describe_schema(value[0], indent+1))
            else:
                lines.append(f"{prefix}- {key}: {value}")
    elif isinstance(schema, list) and schema:
        lines.extend(describe_schema(schema[0], indent))
    return lines

def flatten_dict(d, parent_key='', sep='_'):
    """
    Flatten a nested dictionary into a single-level dictionary.
    """
    items = []
    if isinstance(d, dict):
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(flatten_dict(v, new_key, sep=sep).items())
            elif isinstance(v, list):
                # Convert list to string representation
                items.append((new_key, str(v)))
            else:
                items.append((new_key, v))
    else:
        # Handle non-dict input
        items.append((parent_key or 'value', str(d)))
    return dict(items)

PREDEFINED_SCHEMAS = {
  "Manufacturing Account": {
    "company_name": "Name of the company",
    "report_period": "Period of the manufacturing account",
    "raw_material": {
      "opening_stock": "Opening stock of raw materials",
      "purchases": "Purchases of raw materials",
      "returns_outwards": "Returned raw materials (if any)",
      "closing_stock": "Closing stock of raw materials",
      "cost_of_raw_material_consumed": "Cost of raw materials consumed"
    },
    "direct_labour": {
      "bonus": "Factory bonuses",
      "casual_wages": "Casual wages for factory workers",
      "epf": "EPF contributions",
      "socso": "SOCSO contributions",
      "eis": "EIS contributions",
      "sub_contract_wages": "Sub-contracted wages",
      "wages_salaries": "Wages and salaries",
      "total": "Total direct labour cost"
    },
    "factory_overheads": {
      "depreciation": "Total depreciation for factory assets",
      "factory_expenses": "General factory expenses",
      "total_overheads": "Total factory overheads"
    },
    "total_cost": "Total cost before WIP adjustments",
    "work_in_progress": {
      "opening": "Opening WIP value",
      "closing": "Closing WIP value"
    },
    "production_cost": "Cost of production for the year"
  },
  "Trading P&L": {
    "company_name": "Name of the company",
    "report_period": "Period of the trading and P&L statement",
    "revenue": "Total sales or revenue",
    "cost_of_sales": {
      "opening_stock": "Opening stock of finished goods",
      "manufacturing_cost": "Manufacturing cost carried forward",
      "purchases": "Purchases made (if any)",
      "closing_stock": "Closing stock of finished goods",
      "total": "Total cost of goods sold"
    },
    "gross_profit": "Gross profit from operations",
    "other_income": {
      "interest_income": "Income from interest",
      "misc_income": "Miscellaneous or other income",
      "total": "Total other income (if calculated)"
    },
    "expenses": {
      "administrative": "Administrative and overhead expenses",
      "selling": "Selling expenses",
      "financial": "Financial-related expenses",
      "total": "Total expenses"
    },
    "net_profit": "Net profit before taxation"
  },
  "Balance Sheet": {
    "company_name": "Name of the company",
    "report_date": "Date of the balance sheet",
    "non_current_assets": {
      "property_plant_equipment": "Total fixed assets (PPE)",
      "investments": "Total non-current investments",
      "total": "Total non-current assets"
    },
    "current_assets": {
      "inventory": "Total inventories",
      "trade_receivables": "Receivables including sundry and trade",
      "cash_equivalents": "Cash and bank balances",
      "total": "Total current assets"
    },
    "current_liabilities": {
      "trade_payables": "Payables to suppliers or trade creditors",
      "short_term_loans": "Loans due within one year",
      "total": "Total current liabilities"
    },
    "non_current_liabilities": {
      "long_term_loans": "Loans due after one year",
      "total": "Total non-current liabilities"
    },
    "equity": {
      "share_capital": "Issued share capital",
      "retained_earnings": "Accumulated retained earnings",
      "total": "Total shareholders' equity"
    },
    "total_assets": "Total value of all assets",
    "total_liabilities_equity": "Total liabilities and equity (should balance)"
  },
  "tax_computation": {
    "company_details": {
      "company_name": "Name of the company",
      "tax_file_no": "Tax file number of the company",
      "basis_period": "The basis period for the tax calculation",
      "year_of_assessment": "The year of assessment",
      "ref_no": "Reference number",
      "company_no": "Company registration number",
      "b_code": "Business code"
    },
    "manufacturing_account": {
      "company_name": "Name of the company",
      "report_period": "Period of the manufacturing account",
      "raw_material": {
        "opening_stock": "Opening stock of raw materials",
        "purchases": "Purchases of raw materials",
        "returns_outwards": "Returned raw materials (if any)",
        "closing_stock": "Closing stock of raw materials",
        "cost_of_raw_material_consumed": "Cost of raw materials consumed"
      },
      "direct_labour": {
        "bonus": "Factory bonuses",
        "casual_wages": "Casual wages for factory workers",
        "epf": "EPF contributions",
        "socso": "SOCSO contributions",
        "eis": "EIS contributions",
        "sub_contract_wages": "Sub-contracted wages",
        "wages_salaries": "Wages and salaries",
        "total": "Total direct labour cost"
      },
      "factory_overheads": {
        "depreciation": "Total depreciation for factory assets",
        "factory_expenses": "General factory expenses",
        "total_overheads": "Total factory overheads"
      },
      "total_cost": "Total cost before WIP adjustments",
      "work_in_progress": {
        "opening": "Opening WIP value",
        "closing": "Closing WIP value"
      },
      "production_cost": "Cost of production for the year"
    },
    "business_income_1": {
      "net_profit_per_audited_accounts": "Net profit figure from the audited accounts",
      "non_business_income": {
        "dividend_income": "Total dividend income",
        "rental_income": "Total rental income",
        "interest_income": "Total interest income"
      },
      "non_taxable_gain": {
        "gain_on_disposal_of_fixed_assets": "Gain on the disposal of fixed assets"
      },
      "business_income_loss": "The calculated business income or loss",
      "disallowable_expenses": {
        "interest_restriction_sec33_2": "Interest restriction under Section 33(2)",
        "depreciation_of_building": "Depreciation of building",
        "depreciation_of_leasehold_land": "Depreciation of leasehold land",
        "depreciation_of_building_house": "Depreciation of building - house",
        "depreciation_of_renovation_factory": "Depreciation of renovation factory",
        "depreciation_of_furniture_fittings": "Depreciation of furniture & fittings",
        "depreciation_of_office_equipment": "Depreciation of office equipment",
        "depreciation_of_renovation_house": "Depreciation of renovation - house",
        "depreciation_renovation_office": "Depreciation of renovation - office",
        "depreciation_of_motor_vehicle": "Depreciation of motor vehicle",
        "manufacturing_account_factory_expenses": "Factory expenses from manufacturing account",
        "unrealised_gain_loss_on_foreign_exchange": "Unrealised gain or loss on foreign exchange",
        "donation_subscription": "Donation & subscription",
        "filing_and_attestation_fees": "Filing and attestation fees",
        "depreciation_of_die_mould": "Depreciation of die & mould",
        "general_expenses": "General expenses",
        "depreciation_of_factory_equipment": "Depreciation of factory equipment",
        "gifts": "Gifts",
        "depreciation_of_plant_machinery": "Depreciation of plant & machinery",
        "insurance": "Insurance expenses",
        "legal_fee": "Legal fee",
        "penalty": "Penalty expenses",
        "printing_and_stationery": "Printing and stationery expenses",
        "professional_fee": "Professional fee",
        "quit_rent_assessment": "Quit rent & assessment",
        "refreshment": "Refreshment expenses",
        "audit_fee": "Audit fee",
        "registration_fee": "Registration fee",
        "secretarial_fee": "Secretarial fee",
        "staff_welfare": "Staff welfare expenses",
        "stamp_duties": "Stamp duties",
        "travelling_expenses": "Travelling expenses"
      },
      "allowances_and_deductions": {
        "deduction_of_audit_expenses": "Deduction of audit expenses",
        "deduction_for_secretarial_and_tax_filing_fee": "Deduction for expenses in relation to secretarial fee and tax filing fee"
      },
      "adjusted_income": "The adjusted income after all adjustments",
      "balancing_charge": "Balancing charge",
      "capital_allowance": {
        "balance_brought_forward": "Capital allowance balance brought forward",
        "current_year": "Current year capital allowance",
        "amount_absorbed_this_year": "Capital allowance amount absorbed this year",
        "balance_carried_forward": "Capital allowance balance carried forward"
      },
      "statutory_income": "The final statutory income"
    },
    "aggregate_income": {
      "aggregate_statutory_business_income": "The aggregate statutory business income",
      "other_statutory_income": {
        "dividend_income": {
          "amount": "Dividend income amount",
          "less_direct_expenses": "Less direct expenses"
        },
        "interest_income": {
          "amount": "Interest income amount",
          "less_allowable_interest": "Less allowable interest"
        },
        "deemed_interest_income": {
          "amount": "Deemed interest income amount",
          "less_allowable_interest": "Less allowable interest"
        },
        "rental_income": {
          "amount": "Rental income amount",
          "less_direct_expenses": "Less direct expenses"
        }
      },
      "approved_donations": "Approved donations",
      "total_income": "Total income",
      "chargeable_income": "Chargeable income"
    },
    "tax_payable": {
      "income_tax": {
        "tax_at_15_percent": "Income tax calculated at 15%",
        "tax_at_17_percent": "Income tax calculated at 17%",
        "tax_at_24_percent": "Income tax calculated at 24%"
      },
      "total_tax_payable": "Total tax payable for the year",
      "less_set_off": {
        "section_110_set_off": "Section 110 set off",
        "section_132_133_relief": "Section 132/133 relief"
      },
      "tax_paid": {
        "actual_tax_paid_via_cp204_installment": "Actual tax paid through CP204 installment",
        "tax_over_paid_refundable": "Tax amount that is overpaid and refundable"
      }
    },
    "exempt_account": {
      "exempt_credit_balance_b_f": "Exempt credit balance brought forward",
      "add_items": {
        "incentive_claim": "Incentive claim",
        "exempt_income": "Exempt income"
      },
      "less_items": {
        "current_year_non_pioneer_loss": "Current year non-pioneer loss",
        "exempt_dividend_paid": "Exempt dividend paid / credited / distributed"
      },
      "exempt_credit_balance_c_f": "Exempt credit balance carried forward"
    }
  }

}

# Ensure llm is always defined for test monkeypatching
llm = None  # Will be set in production or monkeypatched in tests

# Pydantic models for API
class ExtractionSchema(BaseModel):
    """Define the schema for data extraction"""
    fields: Dict[str, Any] = Field(description="Field names and their descriptions")
    
class DocumentType(BaseModel):
    """Supported document types"""
    pdf: bool = True
    excel: bool = True
    word: bool = True

class SchemaDetectionResult(BaseModel):
    """Result of schema detection"""
    detected_schema_type: str
    confidence_score: float
    suggested_schema: Dict[str, Any]
    reasoning: str

class ExtractionResult(BaseModel):
    """Structured extraction result"""
    extracted_data: Dict[str, Any]
    file_info: Dict[str, Any]
    processing_time: float
    status: str
    schema_detection: Optional[Union[SchemaDetectionResult, Dict[str, Any], None]] = None
    
    class Config:
        # Allow arbitrary types in model for flexibility
        arbitrary_types_allowed = True
        # Configure JSON serialization to handle custom objects
        json_encoders = {
            # Add custom encoders if needed
        }

# State for LangGraph
class DocumentProcessingState(BaseModel):
    """State for document processing workflow"""
    file_content: Optional[Union[str, bytes]] = None
    file_type: Optional[str] = None
    file_name: Optional[str] = None
    raw_content: Optional[str] = None  # Store raw extracted content
    extraction_schema: Optional[Dict[str, Any]] = None
    detected_schema_type: Optional[str] = None
    schema_confidence: Optional[float] = None
    schema_reasoning: Optional[str] = None
    extracted_data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    processing_stage: str = "initial"
    auto_detect_schema: bool = False
    result_file: Optional[str] = None  # Path to saved result JSON

# Document processing tools

def merge_nested_dicts(d1, d2):
    """Recursively merge two nested dictionaries and lists. Used for aggregating chunked extraction results."""
    for k, v in d2.items():
        if k in d1 and isinstance(d1[k], dict) and isinstance(v, dict):
            merge_nested_dicts(d1[k], v)
        elif k in d1 and isinstance(d1[k], list) and isinstance(v, list):
            d1[k].extend(v)
        else:
            # Only override if the current value is empty or None
            if k not in d1 or d1[k] == "" or d1[k] is None:
                d1[k] = v
    return d1

@tool
def extract_pdf_content(file_path: str) -> str:
    """Extract text content from PDF files (all pages concatenated)"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e2:
            return f"Error extracting PDF: {str(e2)}"

def extract_pdf_pages(file_path: str) -> list:
    """Extract text content from each page of a PDF file as a list of strings (one per page), with OCR fallback for image-based PDFs. Cleans up temp files after OCR."""
    import uuid
    pages = []
    try:
        doc = fitz.open(file_path)
        for i, page in enumerate(doc):
            text = page.get_text()
            if not text or text.strip() == "":
                # OCR fallback for blank/image page
                try:
                    from pdf_image_ocr import pdf_pages_to_images, ocr_images
                    # Create a temp single-page PDF for OCR
                    single_page_temp = os.path.join(tempfile.gettempdir(), f"ocr_page_{os.getpid()}_{uuid.uuid4().hex}.pdf")
                    try:
                        pdf_writer = PyPDF2.PdfWriter()
                        pdf_writer.add_page(PyPDF2.PdfReader(file_path).pages[i])
                        with open(single_page_temp, 'wb') as spf:
                            pdf_writer.write(spf)
                        images = pdf_pages_to_images(single_page_temp, pages=[0])
                        if images:
                            ocr_texts = ocr_images(images)
                            text = ocr_texts[0] if ocr_texts else ""
                            if text:
                                print(f"[PDF OCR] Page {i+1}: Used OCR fallback, extracted {len(text)} chars.")
                            else:
                                print(f"[PDF OCR] Page {i+1}: OCR fallback failed.")
                        else:
                            print(f"[PDF OCR] Page {i+1}: No image generated for OCR fallback.")
                    finally:
                        try:
                            os.remove(single_page_temp)
                        except Exception as ocr_cleanup_err:
                            print(f"[PDF OCR] Failed to delete temp file {single_page_temp}: {ocr_cleanup_err}")
                except Exception as ocr_e:
                    print(f"[PDF OCR] Page {i+1}: OCR extraction error: {ocr_e}")
            pages.append(text)
        doc.close()
        return pages
    except Exception as e:
        print(f"[PDF Extraction] fitz failed: {e}. Trying PyPDF2...")
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if not text or text.strip() == "":
                        # OCR fallback for blank/image page
                        try:
                            from pdf_image_ocr import pdf_pages_to_images, ocr_images
                            images = pdf_pages_to_images(file_path, pages=[i])
                            if images:
                                ocr_texts = ocr_images(images)
                                text = ocr_texts[0] if ocr_texts else ""
                                if text:
                                    print(f"[PDF OCR] Page {i+1}: Used OCR fallback, extracted {len(text)} chars.")
                                else:
                                    print(f"[PDF OCR] Page {i+1}: OCR fallback failed.")
                            else:
                                print(f"[PDF OCR] Page {i+1}: No image generated for OCR fallback.")
                        except Exception as ocr_e:
                            print(f"[PDF OCR] Page {i+1}: OCR extraction error: {ocr_e}")
                    pages.append(text)
                return pages
        except Exception as e2:
            print(f"[PDF Extraction] PyPDF2 also failed: {e2}")
            return []

def ai_document_understanding(pages: list, filename: str) -> dict:
    """
    AI-powered document understanding that analyzes PDF content intelligently
    and creates a comprehensive document summary and key insights.
    """
    global llm
    
    try:
        # Step 1: Create document overview from all pages
        document_overview = create_document_overview(pages, filename)
        
        # Step 2: Identify key sections and data patterns
        key_sections = identify_key_sections(pages)
        
        # Step 3: Extract critical information using AI understanding
        critical_info = extract_critical_information(pages, document_overview)
        
        # Combine all analysis into comprehensive understanding
        analysis_result = {
            "document_overview": document_overview,
            "key_sections": key_sections,
            "critical_information": critical_info,
            "total_pages": len(pages),
            "document_type": determine_document_type_ai(document_overview),
            "confidence_score": calculate_analysis_confidence(document_overview, key_sections, critical_info)
        }
        
        print(f"[AI Understanding] Document analysis completed with confidence: {analysis_result['confidence_score']}")
        return analysis_result
        
    except Exception as e:
        print(f"[AI Understanding] Error in document understanding: {str(e)}")
        return None

def create_document_overview(pages: list, filename: str) -> str:
    """
    Create a comprehensive overview of the document using AI analysis.
    """
    global llm
    
    # Combine first 2 pages and last page for overview (most informative)
    overview_content = ""
    if len(pages) > 0:
        overview_content += f"FIRST PAGE:\n{pages[0][:1500]}\n\n"
    if len(pages) > 1:
        overview_content += f"SECOND PAGE:\n{pages[1][:1500]}\n\n"
    if len(pages) > 2:
        overview_content += f"LAST PAGE:\n{pages[-1][:1500]}\n\n"
    
    prompt = f"""Analyze this document and provide a comprehensive overview. Focus on:
1. Document type and purpose
2. Key entities (companies, people, dates, amounts)
3. Main sections and structure
4. Important data patterns
5. Document context and business purpose

Document: {filename}
Total Pages: {len(pages)}

Content Sample:
{overview_content}

Provide a detailed analysis that captures the essence and structure of this document."""
    
    try:
        messages = [
            {"role": "system", "content": "You are an expert document analyst. Provide comprehensive, structured analysis of documents focusing on business context, data patterns, and key information."},
            {"role": "user", "content": prompt}
        ]
        
        response = llm.invoke(messages)
        content = response.content if hasattr(response, 'content') else str(response)
        return content
        
    except Exception as e:
        print(f"[AI Understanding] Error creating document overview: {str(e)}")
        return f"Document analysis failed: {str(e)}"

def identify_key_sections(pages: list) -> list:
    """
    Identify and categorize key sections across all pages.
    """
    sections = []
    
    for idx, page in enumerate(pages):
        if not page or not page.strip():
            continue
            
        # Look for common section indicators
        page_lower = page.lower()
        section_info = {
            "page_number": idx + 1,
            "content_length": len(page),
            "section_type": "content",
            "key_indicators": []
        }
        
        # Identify section types based on content patterns
        if any(indicator in page_lower for indicator in ['tax', 'computation', 'assessment']):
            section_info["section_type"] = "tax_computation"
            section_info["key_indicators"].append("tax_related")
            
        if any(indicator in page_lower for indicator in ['balance', 'profit', 'loss', 'income']):
            section_info["section_type"] = "financial_statement"
            section_info["key_indicators"].append("financial_data")
            
        if any(indicator in page_lower for indicator in ['table', 'schedule', 'summary']):
            section_info["key_indicators"].append("tabular_data")
            
        if re.search(r'\d+[.,]\d+', page):  # Contains numbers with decimals
            section_info["key_indicators"].append("numerical_data")
            
        sections.append(section_info)
    
    return sections

def extract_critical_information(pages: list, overview: str) -> dict:
    """
    Extract critical information using AI understanding of the document context.
    """
    global llm
    
    # Focus on pages with most content
    important_pages = sorted(enumerate(pages), key=lambda x: len(x[1]), reverse=True)[:3]
    
    critical_content = ""
    for idx, page in important_pages:
        if page and page.strip():
            critical_content += f"\n\nPAGE {idx + 1} (Key Content):\n{page[:2000]}"
    
    prompt = f"""Based on the document overview and key content, extract the most critical information:

Document Overview:
{overview[:1000]}

Key Content:
{critical_content}

Extract and structure:
1. Primary entities (names, IDs, reference numbers)
2. Key financial figures and amounts
3. Important dates and periods
4. Critical business data
5. Document-specific key information

Return as structured text with clear categories."""
    
    try:
        messages = [
            {"role": "system", "content": "You are an expert at extracting critical information from business documents. Focus on the most important data points that define the document's purpose and content."},
            {"role": "user", "content": prompt}
        ]
        
        response = llm.invoke(messages)
        content = response.content if hasattr(response, 'content') else str(response)
        
        # Parse the response into structured format
        return {
            "raw_analysis": content,
            "extraction_method": "ai_understanding",
            "content_analyzed": len(critical_content)
        }
        
    except Exception as e:
        print(f"[AI Understanding] Error extracting critical information: {str(e)}")
        return {"error": str(e), "extraction_method": "failed"}

def determine_document_type_ai(overview: str) -> str:
    """
    Use AI to determine the specific document type based on overview.
    """
    overview_lower = overview.lower()
    
    if any(term in overview_lower for term in ['tax computation', 'assessment', 'tax file']):
        return "tax_computation"
    elif any(term in overview_lower for term in ['balance sheet', 'profit loss', 'financial statement']):
        return "financial_statement"
    elif any(term in overview_lower for term in ['invoice', 'bill', 'receipt']):
        return "invoice"
    elif any(term in overview_lower for term in ['contract', 'agreement', 'terms']):
        return "contract"
    else:
        return "business_document"

def calculate_analysis_confidence(overview: str, sections: list, critical_info: dict) -> float:
    """
    Calculate confidence score for the AI analysis.
    """
    confidence = 0.5  # Base confidence
    
    # Increase confidence based on analysis quality
    if overview and len(overview) > 100:
        confidence += 0.2
    
    if sections and len(sections) > 0:
        confidence += 0.1
        
    if critical_info and "error" not in critical_info:
        confidence += 0.2
    
    return min(confidence, 1.0)

def ai_schema_based_extraction(document_analysis: dict, schema: dict, schema_description: str) -> dict:
    """
    Use AI understanding to apply the detected schema and generate structured JSON output.
    """
    global llm
    
    try:
        # Create intelligent extraction prompt using document analysis
        prompt = f"""Using the comprehensive document analysis provided, extract structured data according to the specified schema.

DOCUMENT ANALYSIS:
Type: {document_analysis.get('document_type', 'unknown')}
Pages: {document_analysis.get('total_pages', 0)}
Confidence: {document_analysis.get('confidence_score', 0)}

Document Overview:
{document_analysis.get('document_overview', '')[:2000]}

Critical Information:
{document_analysis.get('critical_information', {}).get('raw_analysis', '')[:2000]}

SCHEMA TO FOLLOW:
{schema_description}

Based on the document analysis above, extract and structure the data according to the schema. Use the AI understanding of the document to intelligently map content to schema fields. Return valid JSON only."""
        
        messages = [
            {"role": "system", "content": "You are an expert data extraction assistant. Use document analysis to intelligently extract structured data according to schemas. Return only valid JSON."},
            {"role": "user", "content": prompt}
        ]
        
        response = llm.invoke(messages)
        content = response.content if hasattr(response, 'content') else str(response)
        
        # Parse the JSON response
        from robust_json import parse_llm_json
        extracted_data = parse_llm_json(content)
        
        if extracted_data:
            print(f"[AI Extraction] Successfully extracted structured data using AI understanding")
            return extracted_data
        else:
            print(f"[AI Extraction] Failed to parse extracted data")
            return None
            
    except Exception as e:
        print(f"[AI Extraction] Error in schema-based extraction: {str(e)}")
        return None

def extract_pdf_page_by_page_with_schema(state: DocumentProcessingState) -> DocumentProcessingState:
    """
    New implementation: Extract PDF page by page using detected schema,
    save raw text to debug_output, and combine results into Excel file in extraction_results.
    """
    global llm
    
    # Ensure file_type is set
    state.file_type = 'pdf'
    if hasattr(state, 'processing_metadata') and isinstance(state.processing_metadata, dict):
        state.processing_metadata['file_type'] = 'pdf'
    
    # Ensure LLM is initialized
    if llm is None:
        print("LLM is not initialized, attempting to initialize it now")
        try:
            model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("API key is missing or empty. Check your .env file.")
            if "instruct" in model_name:
                from langchain_openai import OpenAI
                llm = OpenAI(
                    model=model_name,
                    temperature=0,
                    openai_api_key=api_key
                )
            else:
                from langchain_openai import ChatOpenAI
                llm = ChatOpenAI(
                    model=model_name,
                    temperature=0,
                    openai_api_key=api_key
                )
            print(f"Initialized LLM: {type(llm).__name__} ({model_name})")
        except Exception as e:
            state.error = f"Failed to initialize LLM: {str(e)}"
            return state
    try:
        # Create temporary file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"pdf_extract_{os.getpid()}_{int(time.time()*1000)}{Path(state.file_name).suffix}")
        
        try:
            with open(temp_path, 'wb') as temp_file:
                if isinstance(state.file_content, bytes):
                    temp_file.write(state.file_content)
                elif isinstance(state.file_content, str):
                    temp_file.write(state.file_content.encode())
                else:
                    state.error = f"Unsupported file_content type: {type(state.file_content)}"
                    return state
        except Exception as file_write_err:
            state.error = f"Failed to write temp PDF: {file_write_err}"
            return state
        
        # Extract pages
        pages = extract_pdf_pages(temp_path)
        if not pages or all((not p or p.strip() == "") for p in pages):
            print(f"[PDF Extraction] All pages empty after extraction for file: {state.file_name}")
            state.error = "No content extracted from PDF"
            return state
        
        print(f"[PDF Processing] Extracted {len(pages)} pages from PDF")
        
        # Create directories
        pdf_raw_dir = os.path.join(os.path.dirname(__file__), "pdf_raw")
        results_dir = os.path.join(os.path.dirname(__file__), "extraction_results")
        os.makedirs(pdf_raw_dir, exist_ok=True)
        os.makedirs(results_dir, exist_ok=True)
        
        # Generate base filename
        base_filename = os.path.splitext(os.path.basename(state.file_name or "pdf_document"))[0]
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Step 1: Save all raw text to pdf_raw folder
        raw_text_file = os.path.join(pdf_raw_dir, f"{base_filename}_raw_{timestamp}.txt")
        all_text = ""
        
        with open(raw_text_file, "w", encoding="utf-8") as f:
            f.write(f"PDF Raw Text Extraction - {state.file_name}\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n")
            f.write(f"Total Pages: {len(pages)}\n")
            f.write("=" * 80 + "\n\n")
            
            for idx, page_text in enumerate(pages):
                page_content = page_text if page_text else "[EMPTY PAGE]"
                f.write(f"PAGE {idx + 1}\n")
                f.write("-" * 40 + "\n")
                f.write(page_content)
                f.write("\n\n" + "=" * 80 + "\n\n")
                
                # Accumulate all text for processing
                if page_text:
                    all_text += f"\n\nPAGE {idx + 1}:\n{page_text}"
        
        print(f"[PDF Processing] Saved raw text to: {raw_text_file}")
        
        # Step 2: Use the raw text file with schema to extract structured data
        if not all_text.strip():
            state.error = "No text content extracted from PDF"
            return state
        
        # Get schema for extraction
        schema = state.extraction_schema or {}
        if not schema:
            print("[WARNING] No schema detected, using default extraction")
            schema = {
                "general_info": "Any general information found",
                "key_data": "Important data points", 
                "tables": "Tabular data if present",
                "financial_data": "Financial information if present"
            }
        
        print(f"[PDF Processing] Using schema: {list(schema.keys())}")
        print(f"[PDF Processing] Processing complete document with {len(pages)} pages")
        
        try:
            # AI Document Understanding Approach
            print(f"[PDF AI Processing] Starting AI document understanding for {len(pages)} pages")
            
            # Step 1: AI Document Analysis and Understanding
            document_analysis = ai_document_understanding(pages, state.file_name)
            
            if not document_analysis:
                print(f"[PDF AI Processing] Document analysis failed")
                extraction_status = "failed"
                extracted_data = {}
            else:
                # Step 2: Apply detected schema to generate structured output
                schema_description = describe_schema(schema)
                extracted_data = ai_schema_based_extraction(document_analysis, schema, schema_description)
                
                if extracted_data:
                    print(f"[PDF AI Processing] Successfully extracted structured data using AI understanding")
                    extraction_status = "success"
                else:
                    print(f"[PDF AI Processing] Failed to extract structured data")
                    extraction_status = "failed"
                    extracted_data = {}
                
        except Exception as e:
            print(f"[PDF Processing] Error processing document: {str(e)}")
            extraction_status = "error"
            extracted_data = {}
        
        # Create JSON file with extracted data
        json_file = os.path.join(results_dir, f"{base_filename}_extracted_{timestamp}.json")
        
        if extracted_data and extraction_status == "success":
            # Save extracted data as JSON
            try:
                with open(json_file, "w", encoding="utf-8") as f:
                    json.dump(extracted_data, f, ensure_ascii=False, indent=2)
                print(f"[PDF Processing] Saved JSON file: {json_file}")
                    
            except Exception as json_error:
                print(f"[PDF Processing] Failed to create JSON file: {json_error}")
                # Fallback: save basic structure
                fallback_data = {
                    'status': f'JSON creation failed: {json_error}',
                    'extracted_data': str(extracted_data) if extracted_data else None
                }
                with open(json_file, "w", encoding="utf-8") as f:
                    json.dump(fallback_data, f, ensure_ascii=False, indent=2)
        else:
            # No data extracted, create summary JSON
            summary_data = {
                'status': f'Extraction {extraction_status}',
                'total_pages': len(pages),
                'file_name': base_filename,
                'extraction_status': extraction_status
            }
            with open(json_file, "w", encoding="utf-8") as f:
                json.dump(summary_data, f, ensure_ascii=False, indent=2)
            print(f"[PDF Processing] Created summary JSON file: {json_file}")
        
        # Create structured output matching the requested format
        final_extracted_data = {
            "_metadata": {
                "pages": [f"Page {i+1}" for i in range(len(pages))],
                "successful_pages": [f"Page {i+1}" for i in range(len(pages))] if extraction_status == "success" else [],
                "page_count": len(pages),
                "successful_count": len(pages) if extraction_status == "success" else 0,
                "extraction_status": extraction_status
            }
        }
        
        # Add the extracted data as "Document" since it's the whole document
        if extracted_data and extraction_status == "success":
            final_extracted_data["Document"] = extracted_data
        
        # Add schema detection info
        schema_detection = {
            "detected_schema_type": "single_schema" if schema else "no_schema",
            "confidence_score": 0.8 if schema and extraction_status == "success" else 0.0,
            "reasoning": f"Schema applied to entire document" if schema else "No schema detected"
        }
        
        # Add processing metadata
        processing_metadata = {
            "file_name": state.file_name or "unknown.pdf",
            "file_type": "pdf",
            "processing_time": time.time() - start_time if 'start_time' in locals() else 0.0,
            "timestamp": datetime.now().isoformat(),
            "debug_files": pdf_raw_dir,
            "result_file": json_file
        }
        
        # Update state with the new structure
        state.extracted_data = {
            "extracted_data": final_extracted_data,
            "schema_detection": schema_detection,
            "processing_metadata": processing_metadata
        }
        state.processing_stage = "data_extracted"
        state.result_file = json_file
        state.raw_content = all_text  # Set for compatibility with schema detection workflow
        
        print(f"[PDF Processing] Processing complete:")
        print(f"  - Raw text saved to: {raw_text_file}")
        print(f"  - JSON results saved to: {json_file}")
        print(f"  - Total pages: {len(pages)}")
        print(f"  - Extraction status: {extraction_status}")
        
        # Clean up temp file
        try:
            os.remove(temp_path)
        except Exception as cleanup_err:
            print(f"[PDF Extraction] Failed to delete temp file {temp_path}: {cleanup_err}")
        
        return state
    except Exception as e:
        state.error = f"PDF page-by-page extraction failed: {str(e)}"
        return state


@tool
def extract_excel_content(file_path: str) -> str:
    """Extract content from Excel files sheet by sheet with efficient chunk processing"""
    try:
        content = ""
        import gc
        
        # First attempt: Use openpyxl in read-only mode for better memory efficiency
        try:
            import openpyxl
            print(f"Processing Excel file: {file_path} with openpyxl")
            
            # Process in chunks to maintain memory efficiency
            CHUNK_SIZE = 1000  # Process this many rows at a time
            
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            sheet_count = len(wb.sheetnames)
            print(f"Found {sheet_count} sheets")
            
            # Process each sheet individually
            for sheet_name in wb.sheetnames:
                try:
                    ws = wb[sheet_name]
                    print(f"Processing sheet: {sheet_name}")
                    content += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    # Get max dimensions with minimal memory impact
                    # Note: This is an estimation for read-only mode
                    max_row = 0
                    max_col = 0
                    
                    # Sample some rows to determine dimensions
                    for row in ws.iter_rows(min_row=1, max_row=10):
                        max_row += 1
                        row_len = len([cell for cell in row if cell.value is not None])
                        max_col = max(max_col, row_len)
                    
                    # Continue counting rows - count up to MAX_SAMPLE_ROWS to prevent memory issues
                    MAX_SAMPLE_ROWS = 200  # Increase this for larger documents
                    empty_row_count = 0
                    consecutive_empty_rows_limit = 10  # Only stop if we see many empty rows in a row
                    
                    for row in ws.iter_rows(min_row=11, max_row=MAX_SAMPLE_ROWS):
                        if any(cell.value is not None for cell in row):
                            max_row += 1
                            empty_row_count = 0  # Reset empty row counter
                        else:
                            # Count empty rows but don't immediately stop
                            empty_row_count += 1
                            
                            # Still count empty rows as rows
                            max_row += 1
                            
                            # Only stop if we have many consecutive empty rows
                            # This preserves empty rows between sections
                            if empty_row_count > consecutive_empty_rows_limit:
                                break
                    
                    print(f"Detected dimensions: ~{max_row} rows × ~{max_col} columns")
                    
                    # Process in chunks to maintain memory efficiency
                    total_rows_processed = 0
                    chunk_start = 1  # 1-based indexing for Excel
                    
                    while chunk_start <= max_row:
                        chunk_end = min(chunk_start + CHUNK_SIZE - 1, max_row)
                        print(f"Processing rows {chunk_start}-{chunk_end}")
                        
                        # Process this chunk
                        consecutive_empty = 0
                        max_consecutive_empty = 3  # Keep at most 3 consecutive empty rows in output
                        
                        for row_idx, row in enumerate(ws.iter_rows(min_row=chunk_start, max_row=chunk_end, values_only=True)):
                            # Handle empty rows differently - include some but not all consecutive empty rows
                            if not any(cell for cell in row):
                                consecutive_empty += 1
                                # Only include a limited number of consecutive empty rows
                                if consecutive_empty > max_consecutive_empty:
                                    continue  # Skip excess consecutive empty rows
                                # Add an empty line to preserve document structure
                                content += "\n"
                            else:
                                consecutive_empty = 0  # Reset empty row counter on non-empty row
                                
                                # Format as tab-separated values
                                row_text = "  ".join(str(cell if cell is not None else "") for cell in row)
                                content += row_text + "\n"
                            
                            total_rows_processed += 1
                        
                        # Move to next chunk
                        chunk_start = chunk_end + 1
                        
                        # Force garbage collection between chunks
                        gc.collect()
                    
                    content += f"\n--- Processed {total_rows_processed} rows ---\n\n"
                    
                except Exception as sheet_error:
                    content += f"\n--- Sheet: {sheet_name} --- (Error: {str(sheet_error)})\n\n"
            
            # Close workbook explicitly
            wb.close()
            
        except Exception as e:
            # Fallback to pandas for problematic Excel files
            print(f"Openpyxl failed: {str(e)}. Falling back to pandas.")
            content += f"\nFalling back to alternative Excel reader due to error: {str(e)}\n"
            
            import pandas as pd
            
            # Process in chunks using pandas
            try:
                # Get sheet names
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    content += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    # First get row count
                    try:
                        # Peek at the sheet to get dimensions
                        df_info = pd.read_excel(file_path, sheet_name=sheet_name, nrows=0)
                        print(f"Sheet {sheet_name}: Determined {len(df_info.columns)} columns")
                        
                        # Process in chunks of 1000 rows
                        chunk_size = 1000
                        chunk_start = 0
                        total_rows = 0
                        
                        while True:
                            # Read this chunk
                            df_chunk = pd.read_excel(
                                file_path, 
                                sheet_name=sheet_name, 
                                skiprows=chunk_start,
                                nrows=chunk_size
                            )
                            
                            # Stop if no data
                            if df_chunk.empty:
                                break
                                
                            # Add to content
                            if chunk_start == 0:  # First chunk includes headers
                                content += df_chunk.to_string(index=False)
                            else:  # Later chunks skip header
                                content += df_chunk.to_string(index=False, header=False)
                                
                            content += "\n"
                            
                            # Track progress
                            rows_in_chunk = len(df_chunk)
                            total_rows += rows_in_chunk
                            print(f"Processed {rows_in_chunk} rows from sheet {sheet_name}, total: {total_rows}")
                            
                            # Move to next chunk
                            chunk_start += rows_in_chunk
                            
                            # If this chunk wasn't full, we've reached the end
                            if rows_in_chunk < chunk_size:
                                break
                                
                            # Free memory
                            del df_chunk
                            gc.collect()
                            
                        content += f"\n--- Processed {total_rows} rows ---\n\n"
                        
                    except Exception as chunk_error:
                        content += f"\nError processing sheet {sheet_name}: {str(chunk_error)}\n"
                
            except Exception as pd_error:
                content += f"\nFailed to process Excel file: {str(pd_error)}\n"
                
        return content
    except Exception as e:
        return f"Error extracting Excel: {str(e)}"

@tool
def extract_word_content(file_path: str) -> str:
    """Extract content from Word documents"""
    try:
        doc = docx.Document(file_path)
        content = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
            
        # Extract tables
        for table in doc.tables:
            content += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                content += " | ".join(row_text) + "\n"
            content += "\n"
            
        return content
    except Exception as e:
        return f"Error extracting Word: {str(e)}"

# LangGraph workflow nodes
def determine_file_type(state: DocumentProcessingState) -> DocumentProcessingState:
    """Determine the file type based on extension"""
    if state.file_name:
        extension = Path(state.file_name).suffix.lower()
        if extension == '.pdf':
            state.file_type = 'pdf'
        elif extension in ['.xlsx', '.xls']:
            state.file_type = 'excel'
        elif extension in ['.docx', '.doc']:
            state.file_type = 'word'
        else:
            state.error = f"Unsupported file type: {extension}"
    
    state.processing_stage = "file_type_determined"
    return state

def extract_document_content(state: DocumentProcessingState) -> DocumentProcessingState:
    """Extract content based on file type"""
    if state.error:
        return state
        
    try:
        # Create temporary file with proper handling for bytes
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(state.file_name).suffix) as temp_file:
            if isinstance(state.file_content, bytes):
                temp_file.write(state.file_content)
            elif isinstance(state.file_content, str):
                temp_file.write(state.file_content.encode())
            else:
                state.error = f"Unsupported file_content type: {type(state.file_content)}"
                return state
            temp_path = temp_file.name
        
        # Extract based on file type
        if state.file_type == 'pdf':
            # Use new schema-based page-by-page PDF extraction
            print("INFO: Using new schema-based PDF page-by-page extraction workflow.")
            return extract_pdf_page_by_page_with_schema(state)

        elif state.file_type == 'excel':
            if extract_excel_content is None:
                state.error = "Excel extraction function is not initialized."
                return state
                
            # Extract content from Excel file
            if hasattr(extract_excel_content, 'invoke'):
                content = extract_excel_content.invoke({"file_path": temp_path})
            else:
                content = extract_excel_content(temp_path)
                
            # Debug logging for Excel content
            content_length = len(content) if content else 0
            sample_size = min(500, content_length)  # First 500 chars
            print(f"DEBUG: Extracted Excel content length: {content_length} characters")
            print(f"DEBUG: Content sample: {content[:sample_size]}...")
            
            # Write extracted content to debug file
            debug_dir = "debug_output"
            os.makedirs(debug_dir, exist_ok=True)
            debug_file = os.path.join(debug_dir, f"{os.path.basename(state.file_name or 'excel')}_raw.txt")
            with open(debug_file, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"DEBUG: Saved raw content to {debug_file}")
        elif state.file_type == 'word':
            if extract_word_content is None:
                state.error = "Word extraction function is not initialized."
                return state
            if hasattr(extract_word_content, 'invoke'):
                content = extract_word_content.invoke({"file_path": temp_path})
            else:
                content = extract_word_content(temp_path)
        else:
            state.error = f"Unsupported file type: {state.file_type}"
            return state
        
        # Store both raw content and processed content
        # (PDF processing is handled entirely within extract_pdf_page_by_page_with_schema)
        if state.file_type != 'pdf':
            state.raw_content = content
            state.processing_stage = "content_extracted"
            # Clean up temporary file (PDF function handles its own cleanup)
            try:
                os.unlink(temp_path)
            except Exception as cleanup_err:
                print(f"Warning: Failed to cleanup temp file {temp_path}: {cleanup_err}")

    except Exception as e:
        state.error = f"Error extracting content: {str(e)}"
    
    return state

def detect_schema(state: DocumentProcessingState) -> DocumentProcessingState:
    """AI agent to detect the most appropriate schema for the document"""
    # Import at function level to ensure fresh import
    from langchain_openai import OpenAI, ChatOpenAI
    import os
    
    # Re-initialize LLM locally to avoid scope issues
    try:
        # Get configuration
        model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
        api_key = os.getenv("OPENAI_API_KEY")
        
        print(f"Schema detection - Using model: {model_name}")
        
        # Initialize LLM based on model type
        if "instruct" in model_name:
            local_llm = OpenAI(
                model=model_name,
                temperature=0,
                openai_api_key=api_key
            )
            print("Schema detection - Successfully initialized OpenAI completion model")
        else:
            local_llm = ChatOpenAI(
                model=model_name,
                temperature=0,
                openai_api_key=api_key
            )
            print("Schema detection - Successfully initialized ChatOpenAI model")
            
        print(f"Schema detection - LLM type: {type(local_llm).__name__}")
        
    except Exception as e:
        print(f"Schema detection - Error initializing LLM: {str(e)}")
        state.error = f"Schema detection failed: Error initializing LLM: {str(e)}"
        return state
    if state.error or not state.auto_detect_schema:
        return state
    
    try:
        # Prepare schema descriptions for the AI
        schema_descriptions = {}
        for schema_name, schema_fields in PREDEFINED_SCHEMAS.items():
            schema_descriptions[schema_name] = {
                "description": f"Schema for {schema_name.replace('_', ' ').title()}",
                "fields": list(schema_fields.keys())
            }
        
        # Create detection prompt
        prompt = f"""
        Analyze the following document content and determine which extraction schema would be most appropriate.
        
        Available schemas:
        {json.dumps(schema_descriptions, indent=2)}
        
        Document content (first 2000 characters):
        {state.raw_content[:2000]}
        
        Please respond with a JSON object containing:
        1. "detected_schema": The name of the most appropriate schema from the available options
        2. "confidence": A confidence score between 0.0 and 1.0
        3. "reasoning": A brief explanation of why this schema was chosen
        
        Consider the following factors:
        - Document structure and layout
        - Key terms and vocabulary used
        - Presence of specific data fields
        - Overall document purpose and context
        
        If none of the predefined schemas fit well, choose "general" as the fallback.
        
        Example response:
        {{
            "detected_schema": "financial_report",
            "confidence": 0.85,
            "reasoning": "Document contains financial terms like revenue, profit, assets, and appears to be a company's annual report."
        }}
        """
        
        # Use LLM for schema detection
        messages = [
            SystemMessage(content="You are an expert document analyst specialized in classifying document types and determining appropriate extraction schemas. Always return valid JSON."),
            HumanMessage(content=prompt)
        ]
        
        # Get response from LLM
        if isinstance(local_llm, ChatOpenAI):
            print("Schema detection - Using ChatOpenAI model for schema detection")
            response = local_llm.invoke(messages)
            response_text = response.content
        else:
            print("Schema detection - Using OpenAI Completion model for schema detection")
            response = local_llm.invoke(prompt)
            response_text = response
        
        # Parse detection result
        try:
            print(f"[DEBUG] Raw LLM schema detection output: {response_text}")
            detection_result = json.loads(response_text)
            detected_schema_name = detection_result.get("detected_schema")
            confidence = detection_result.get("confidence", 0.0)
            reasoning = detection_result.get("reasoning", "No reasoning provided")

            # Defensive: Fill with safe defaults if any field is missing or None
            if not detected_schema_name or detected_schema_name not in PREDEFINED_SCHEMAS:
                print(f"[WARNING] LLM returned invalid or missing detected_schema: {detected_schema_name}. Using fallback 'general'.")
                detected_schema_name = "general"
                confidence = 0.0
                reasoning = "LLM returned invalid or empty detected_schema."
            if confidence is None or not isinstance(confidence, (float, int)):
                print(f"[WARNING] LLM returned invalid confidence: {confidence}. Using 0.0.")
                confidence = 0.0
            if not reasoning or not isinstance(reasoning, str):
                print(f"[WARNING] LLM returned invalid reasoning: {reasoning}. Using fallback.")
                reasoning = "LLM returned invalid or empty reasoning."

            # Update state with detection results
            state.detected_schema_type = detected_schema_name
            state.schema_confidence = confidence
            state.schema_reasoning = reasoning
            state.extraction_schema = PREDEFINED_SCHEMAS[detected_schema_name]
            state.processing_stage = "schema_detected"
            
        except json.JSONDecodeError:
            print(f"[ERROR] JSONDecodeError in schema detection. Raw LLM output: {response_text}")
            state.detected_schema_type = "general"
            state.schema_confidence = 0.0
            state.schema_reasoning = "LLM returned invalid JSON."
            state.extraction_schema = PREDEFINED_SCHEMAS["general"]
            state.processing_stage = "schema_detected"
        except Exception as e:
            print(f"[ERROR] Unexpected error in schema detection: {e}. Raw LLM output: {response_text}")
            state.detected_schema_type = "general"
            state.schema_confidence = 0.0
            state.schema_reasoning = f"Unexpected error: {e}"
            state.extraction_schema = PREDEFINED_SCHEMAS["general"]
            state.processing_stage = "schema_detected"
        
    except Exception as e:
        state.error = f"Schema detection failed: {str(e)}"
        return state
        
    return state
    

# Removed partial data extraction function

def handle_large_pdf_extraction(state: DocumentProcessingState) -> DocumentProcessingState:
    print("Entered handle_large_pdf_extraction")
    """
    Handle extraction for large PDFs by splitting into chunks (pages),
    preprocessing, summarizing, and schema-aware enhancement before LLM extraction.
    Mirrors the Excel abstraction pipeline for robustness and efficiency.
    """
    global llm
    import os
    from langchain.schema import SystemMessage, HumanMessage
    from content_trimmer import trim_content_for_model
    from pdf_preprocessing import enhance_pdf_extraction

    pdf_chunks = state.pdf_chunks
    print(f"PDF chunk count: {len(pdf_chunks)}")
    chunk_names = [f"Page {i+1}" for i in range(len(pdf_chunks))]
    chunk_schemas = {}
    extraction_results = []

    # First pass: schema detection for each chunk (same as before)
    for i, (chunk_name, chunk_content) in enumerate(zip(chunk_names, pdf_chunks)):
        print(f"Detecting schema for PDF chunk {i+1}/{len(pdf_chunks)}: {chunk_name}")
        content_preview = chunk_content[:2000]
        schema_detection_prompt = f"""
        TASK: Identify the appropriate schema for the following PDF page:
        
        Page: {chunk_name}
        
        Content Preview:
        {content_preview}
        
        Based on the content, determine the most appropriate schema from these options:
        
        1. MANUFACTURING_ACCOUNT: Contains raw materials, direct labor, factory overheads, work-in-progress.
        2. PROFIT_AND_LOSS: Contains revenue, expenses, gross profit, net profit.
        3. BALANCE_SHEET: Contains assets, liabilities, equity.
        4. TAX_COMPUTATION: Contains details of tax calculation, statutory income, and allowances.

        Respond with just the schema type (e.g., "MANUFACTURING_ACCOUNT", "PROFIT_AND_LOSS", "BALANCE_SHEET", "TAX_COMPUTATION") and nothing else.
        """
        try:
            schema_response = llm.invoke([HumanMessage(content=schema_detection_prompt)])
            schema_content = schema_response.content if hasattr(schema_response, 'content') else schema_response
            schema_type = schema_content.strip().replace('"', '').upper()
            if 'MANUFACTURING' in schema_type:
                chunk_schemas[chunk_name] = "MANUFACTURING_ACCOUNT"
            elif any(term in schema_type for term in ['P&L', 'PROFIT', 'LOSS']):
                chunk_schemas[chunk_name] = "PROFIT_AND_LOSS"
            elif 'BALANCE' in schema_type:
                chunk_schemas[chunk_name] = "BALANCE_SHEET"
            elif 'TAX' in schema_type:
                chunk_schemas[chunk_name] = "TAX_COMPUTATION"
            else:
                print(f"Unrecognized schema type for {chunk_name}: {schema_type}, using GENERIC")
                chunk_schemas[chunk_name] = "GENERIC"
            print(f"Detected schema for {chunk_name}: {chunk_schemas[chunk_name]}")
        except Exception as e:
            print(f"Error detecting schema for {chunk_name}: {str(e)}. Using GENERIC.")
            chunk_schemas[chunk_name] = "GENERIC"

    # Second pass: extract data using enhanced PDF abstraction
    merged_result = {}
    for i, (chunk_name, chunk_content) in enumerate(zip(chunk_names, pdf_chunks)):
        print(f"Enhancing and extracting data for PDF chunk {i+1}/{len(pdf_chunks)}: {chunk_name}")
        print(f"Processing PDF chunk: {chunk_name}")
        schema_type = chunk_schemas.get(chunk_name, "GENERIC")
        # Build chunk-specific schema (mirror Excel logic)
        if schema_type == "MANUFACTURING_ACCOUNT":
            chunk_schema = {
                "company_name": "Name of the company",
                "report_period": "Period of the manufacturing account",
                "raw_material": {
                    "opening_stock": "Opening stock of raw materials",
                    "purchases": "Raw materials purchased",
                    "returns_outwards": "Returned raw materials",
                    "closing_stock": "Closing stock of raw materials",
                    "cost_of_raw_material_consumed": "Total cost of raw material consumed"
                },
                "direct_labour": {
                    "bonus": "Factory bonus",
                    "casual_wages": "Casual factory wages",
                    "epf": "EPF for factory staff",
                    "socso": "SOCSO for factory staff",
                    "eis": "EIS for factory staff",
                    "sub_contract_wages": "Subcontractor wages",
                    "wages_salaries": "Factory wages and salaries",
                    "total": "Total direct labour cost"
                },
                "factory_overheads": {
                    "depreciation": "Depreciation expenses",
                    "factory_expenses": "General factory expenses",
                    "total_overheads": "Total factory overheads"
                },
                "total_cost": "Total manufacturing cost",
                "work_in_progress": {
                    "opening": "Opening work in progress",
                    "closing": "Closing work in progress"
                },
                "production_cost": "Cost of production transferred to trading account"
            }
        elif schema_type == "PROFIT_AND_LOSS":
            chunk_schema = {
                "company_name": "Name of the company",
                "report_period": "Period of the profit and loss statement",
                "revenue": "Sales or revenue amount",
                "cost_of_sales": {
                    "opening_stock": "Opening finished goods stock",
                    "manufacturing_cost": "Cost of manufacturing",
                    "purchases": "Purchases",
                    "closing_stock": "Closing finished goods stock",
                    "total": "Total cost of sales"
                },
                "gross_profit": "Gross profit",
                "other_income": "Other income",
                "expenses": {},
                "total_expenses": "Total expenses",
                "operating_profit": "Operating profit",
                "net_profit": "Net profit"
            }
        elif schema_type == "BALANCE_SHEET":
            chunk_schema = PREDEFINED_SCHEMAS.get("Balance Sheet", {})
        elif schema_type == "TAX_COMPUTATION":
            chunk_schema = PREDEFINED_SCHEMAS.get("tax_computation", {})
        else:
            chunk_schema = {}

        # --- PDF abstraction enhancement step ---
        enhanced_content = enhance_pdf_extraction(chunk_content, chunk_schema)

        # Now use the enhanced_content for LLM extraction
        try:
            prompt = f"""Extract structured data from the following PDF page chunk (Page {i+1}) based on the provided schema. Respond with JSON only.

Schema:
{json.dumps(chunk_schema, indent=2)}

Content:
{enhanced_content}"""
            response = llm.invoke([SystemMessage(content="You are an expert in financial document extraction."), HumanMessage(content=prompt)])
            content = response.content if hasattr(response, 'content') else response
            # Try to parse as JSON
            import json
            # Save raw LLM output for debugging (PDF extraction)
            debug_output_path = os.path.join('debug_output', f"{chunk_name.replace(' ', '_')}_raw.txt")
            os.makedirs('debug_output', exist_ok=True)
            with open(debug_output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            try:
                extracted = json.loads(content)
            except Exception as parse_e:
                print(f"Failed to parse LLM output as JSON for chunk {chunk_name}: {parse_e}")
                extracted = {"_raw": content}
            extraction_results.append({chunk_name: extracted})
            # Optionally, merge into merged_result
            if isinstance(extracted, dict):
                merged_result.update(extracted)
        except Exception as e:
            print(f"Error extracting data for {chunk_name}: {e}")
            extraction_results.append({chunk_name: {"_error": str(e)}})

    # Store results in state
    state.extracted_data = merged_result
    state.processing_stage = "data_extracted"
    return state

def extract_with_simplified_prompt(state: DocumentProcessingState) -> DocumentProcessingState:
    """Extract data using a simplified prompt to fit within token limits
    
    Args:
        state: The current document processing state
        
    Returns:
        Updated DocumentProcessingState with extracted data
    """
    global llm
    print("Using simplified prompt for large document")
    
    # Create a simplified system message
    system_msg = """You are a document information extraction expert working with a large document.
    Extract ONLY the key information that matches the requested fields.
    Be concise and precise - only extract what is explicitly in the document.
    Format all financial values as strings to preserve formatting.
    Return a valid JSON object matching the provided schema."""
    
    # Create a simplified extraction prompt
    schema_description = describe_schema(state.extraction_schema)
    
    # Dynamically trim content to fit model context window
    from tiktoken_util import count_tokens
    model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo-instruct")
    model_max_tokens = 4097 if "gpt-3.5" in model_name else 8192  # Adjust for other models if needed
    max_completion_tokens = 512  # Set a safe completion size

    # Build the prompt step by step, trimming as needed
    prompt_template = f"""
    TASK: Extract information from this document according to this schema:
    {schema_description}

    DOCUMENT PREVIEW (may be truncated due to size):
    {{content_preview}}

    Format your response as a valid JSON object with ONLY the extracted field values.
    If a field is not found in the document, include it with an empty string value.
    """
    content_preview = state.file_content
    system_tokens = count_tokens(system_msg, model_name)
    # Iteratively trim content_preview until the total tokens fit
    while True:
        extraction_prompt = prompt_template.format(content_preview=content_preview)
        prompt_tokens = count_tokens(extraction_prompt, model_name)
        total_tokens = system_tokens + prompt_tokens + max_completion_tokens
        if total_tokens <= model_max_tokens or len(content_preview) < 1000:
            break
        # Trim by 10% each time if too large
        content_preview = content_preview[:int(len(content_preview) * 0.9)]
    if total_tokens > model_max_tokens:
        # Still too large, issue a warning and use a minimal preview
        print(f"WARNING: Document preview still too large after trimming. Forcing minimal preview.")
        content_preview = content_preview[:2000]
        extraction_prompt = prompt_template.format(content_preview=content_preview)
    if len(content_preview) < len(state.file_content):
        print(f"Trimmed document preview to {len(content_preview)} chars to fit model context window.")
    # Run extraction through LLM
    messages = [
        SystemMessage(content=system_msg),
        HumanMessage(content=extraction_prompt)
    ]
    try:
        # Pass max_tokens if supported by your LLM interface
        response = llm.invoke(messages, max_tokens=max_completion_tokens) if hasattr(llm, 'invoke') and 'max_tokens' in llm.invoke.__code__.co_varnames else llm.invoke(messages)
        content = response.content if hasattr(response, 'content') else response
        try:
            extracted_data = json.loads(content)
            state.extracted_data = extracted_data
            state.processing_stage = "data_extracted"
            print("Successfully extracted data with simplified prompt")
        except json.JSONDecodeError as e:
            print(f"Failed to parse JSON response: {str(e)}")
            state.error = f"Failed to parse JSON response: {str(e)}"
    except Exception as e:
        state.error = f"Error in extraction: {str(e)}"
    return state


def extract_structured_data(state: DocumentProcessingState) -> DocumentProcessingState:
    """AI-powered structured data extraction using document understanding and schema-based extraction."""
    print(f"Starting AI-powered structured data extraction for: {state.file_name}")
    
    # Check if PDF data has already been extracted (skip if already processed)
    if state.file_type == 'pdf' and state.processing_stage == "data_extracted" and state.extracted_data:
        print(f"[PDF Processing] Data already extracted, skipping extract_structured_data step")
        return state
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    start_time = datetime.now()
    
    # Create debug directory
    debug_dir = "debug_output"
    os.makedirs(debug_dir, exist_ok=True)
    
    # Access the global LLM instance
    global llm
    
    # Check if LLM is initialized
    if llm is None:
        print("LLM is not initialized, attempting to initialize it now")
        try:
            # Check if the specified model is a completion model or a chat model
            model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
            # Set model context window sizes based on known limits
            context_window = 4097  # Default for gpt-3.5-turbo-instruct
            if "gpt-4" in model_name:
                context_window = 8192  # Default for gpt-4
            if "gpt-4-turbo" in model_name:
                context_window = 128000  # For gpt-4-turbo
            if "gpt-3.5-turbo-16k" in model_name:
                context_window = 16384  # For 16k variant
            api_key = os.getenv("OPENAI_API_KEY")
            
            if not api_key:
                raise ValueError("API key is missing or empty. Check your .env file.")
                
            if "instruct" in model_name:
                from langchain_openai import OpenAI
                # Reserve 25% for prompt, 75% for completion (with a reasonable default)
                max_completion_tokens = min(int(context_window * 0.75), 2500)
                
                llm = OpenAI(
                    model=model_name,
                    temperature=0,
                    openai_api_key=api_key,
                    max_tokens=max_completion_tokens
                )
                print(f"Initialized OpenAI completion model: {model_name} with max_tokens={max_completion_tokens} (context window: {context_window})")
            else:
                # Default to ChatOpenAI for other models
                from langchain_openai import ChatOpenAI
                # Reserve 25% for prompt, 75% for completion (with a reasonable default)
                max_completion_tokens = min(int(context_window * 0.75), 2500)
                
                llm = ChatOpenAI(
                    model=model_name,
                    temperature=0,
                    openai_api_key=api_key,
                    max_tokens=max_completion_tokens
                )
                print(f"Initialized ChatOpenAI model: {model_name} with max_tokens={max_completion_tokens} (context window: {context_window})")
        except Exception as e:
            state.error = f"Failed to initialize LLM: {str(e)}"
            return state
    
    # Verify LLM is properly initialized before using it
    if llm is None:
        state.error = "LLM initialization failed. Cannot proceed with extraction."
        return state
        
    try:
        # --- PDF Chunked Extraction (Excel-style: schema per chunk) ---
        if hasattr(state, 'pdf_chunks') and state.pdf_chunks:
            print(f"PDF chunked extraction: {len(state.pdf_chunks)} chunks detected.")
            return handle_large_pdf_extraction(state)

        if not state.file_content or not state.extraction_schema:
            state.error = "Missing file content or schema for extraction"
            return state
            
        # Step 1: AI Document Understanding (different for Excel vs PDF)
        print("Step 1: Performing AI document understanding...")
        
        try:
            if state.file_type == 'excel':
                # For Excel files, parse the content into sheets data
                sheets_data = {}
                if state.raw_content:
                    # Parse the raw content to extract sheet information
                    current_sheet = None
                    sheet_content = []
                    
                    for line in state.raw_content.split('\n'):
                        if line.startswith('--- Sheet:') and line.endswith('---'):
                            # Save previous sheet if exists
                            if current_sheet and sheet_content:
                                sheets_data[current_sheet] = '\n'.join(sheet_content)
                            # Start new sheet
                            current_sheet = line.replace('--- Sheet:', '').replace('---', '').strip()
                            sheet_content = []
                        else:
                            sheet_content.append(line)
                    
                    # Save last sheet
                    if current_sheet and sheet_content:
                        sheets_data[current_sheet] = '\n'.join(sheet_content)
                
                # Use Excel-specific AI understanding
                document_analysis = ai_excel_document_understanding(sheets_data, state.file_name)
                print(f"Excel document analysis completed. Document type: {document_analysis.get('document_type', 'Unknown')}")
            else:
                # For PDF files, use the existing function
                document_content = {
                    "content": state.file_content,
                    "filename": state.file_name,
                    "file_type": state.file_type
                }
                document_analysis = ai_document_understanding(document_content, state.file_name)
                print(f"Document analysis completed. Document type: {document_analysis.get('document_type', 'Unknown')}")
            
            # Save document analysis for debugging
            analysis_debug_path = os.path.join(debug_dir, f"{timestamp}_document_analysis.json")
            with open(analysis_debug_path, "w", encoding="utf-8") as f:
                json.dump(document_analysis, f, indent=2, ensure_ascii=False)
            print(f"Document analysis saved to: {analysis_debug_path}")
            
        except Exception as e:
            print(f"Warning: AI document understanding failed: {str(e)}")
            # Fallback to basic analysis
            document_analysis = {
                "document_type": "Unknown",
                "confidence": 0.5,
                "key_sections": [],
                "data_structure": "unstructured",
                "analysis_notes": f"AI analysis failed: {str(e)}"
            }
        
        # Step 2: AI Schema-Based Extraction (different for Excel vs PDF)
        print("Step 2: Performing AI schema-based extraction...")
        try:
            if state.file_type == 'excel':
                # Use Excel-specific schema-based extraction
                schema_description = describe_schema(state.extraction_schema)
                extracted_data = ai_excel_schema_based_extraction(
                    document_analysis, 
                    state.extraction_schema, 
                    schema_description,
                    sheets_data
                )
            else:
                # Use PDF schema-based extraction
                extracted_data = ai_schema_based_extraction(
                    document_content, 
                    state.extraction_schema, 
                    document_analysis
                )
            
            # Validate extraction result
            if extracted_data and isinstance(extracted_data, dict):
                # Add metadata about the extraction process
                processing_time = (datetime.now() - start_time).total_seconds()
                
                # Structure the result based on file type
                if state.file_type == 'excel':
                    # For Excel files, create the expected structure with sheets metadata
                    sheet_names = list(sheets_data.keys()) if 'sheets_data' in locals() else []
                    
                    # Create the proper Excel structure
                    final_result = {
                        "extracted_data": {
                            "_metadata": {
                                "sheets": sheet_names,
                                "successful_sheets": sheet_names,  # Assume all sheets processed successfully
                                "sheet_count": len(sheet_names),
                                "successful_count": len(sheet_names),
                                "extraction_status": "success"
                            }
                        },
                        "schema_detection": {
                            "detected_schema_type": state.detected_schema_type or "multiple_schemas",
                            "confidence_score": state.schema_confidence or 0.9,
                            "reasoning": state.schema_reasoning or "Multiple schemas detected for different sheets"
                        },
                        "processing_metadata": {
                            "file_name": state.file_name,
                            "file_type": "excel",
                            "processing_time": processing_time,
                            "timestamp": datetime.now().isoformat(),
                            "debug_files": debug_dir,
                            "ai_powered": True,
                            "extraction_method": "ai_document_understanding"
                        }
                    }
                    
                    # Add the actual extracted data for each sheet
                    # The extracted_data from Excel AI contains sheet-specific data
                    if extracted_data and isinstance(extracted_data, dict):
                        # The extracted_data is a dictionary with sheet names as keys
                        # Each sheet contains the structured financial data
                        for sheet_name in sheet_names:
                            if sheet_name in extracted_data:
                                # Add the sheet data directly to the final result
                                final_result["extracted_data"][sheet_name] = extracted_data[sheet_name]
                                print(f"Added data for sheet: {sheet_name}")
                            else:
                                # If no specific sheet data found, log it
                                print(f"No data found for sheet: {sheet_name}")
                                final_result["extracted_data"][sheet_name] = {}
                        
                        print(f"Excel extraction completed. Sheet data keys: {list(extracted_data.keys())}")
                        print(f"Expected sheet names: {sheet_names}")
                    
                    state.extracted_data = final_result
                else:
                    # For PDF and other files, use the original structure
                    extraction_metadata = {
                        "extraction_method": "ai_powered",
                        "document_analysis": document_analysis,
                        "processing_time_seconds": processing_time,
                        "timestamp": timestamp,
                        "file_type": state.file_type,
                        "original_filename": state.file_name
                    }
                    
                    # Add metadata to the extracted data
                    if "_metadata" not in extracted_data:
                        extracted_data["_metadata"] = {}
                    extracted_data["_metadata"].update(extraction_metadata)
                    
                    state.extracted_data = extracted_data
                state.processing_stage = "data_extracted"
                
                print(f"AI-powered extraction completed successfully in {processing_time:.2f} seconds")
                print(f"Extracted {len([k for k in extracted_data.keys() if not k.startswith('_')])} main data fields")
                
                # Save extraction result for debugging
                result_debug_path = os.path.join(debug_dir, f"{timestamp}_extraction_result.json")
                with open(result_debug_path, "w", encoding="utf-8") as f:
                    json.dump(extracted_data, f, indent=2, ensure_ascii=False)
                print(f"Extraction result saved to: {result_debug_path}")
                
            else:
                raise Exception("AI extraction returned invalid or empty data")
                
        except Exception as e:
            print(f"AI schema-based extraction failed: {str(e)}")
            state.error = f"AI-powered extraction failed: {str(e)}"
            return state
            
    except Exception as e:
        print(f"Error in AI-powered extraction: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        state.error = f"Error extracting structured data: {str(e)}"
    
    return state


def finalize_extraction(state: DocumentProcessingState) -> DocumentProcessingState:
    """Finalize the extraction process"""
    if not state.error and state.extracted_data:
        # Ensure all schema fields are present
        for field in state.extraction_schema.keys():
            if field not in state.extracted_data:
                state.extracted_data[field] = ""
        
        state.processing_stage = "completed"
    
    return state

# Create LangGraph workflow
def create_extraction_workflow():
    """Create the document extraction workflow with schema detection"""
    workflow = StateGraph(DocumentProcessingState)
    
    # Add nodes
    workflow.add_node("determine_file_type", determine_file_type)
    workflow.add_node("extract_content", extract_document_content)
    workflow.add_node("detect_schema", detect_schema)
    workflow.add_node("extract_data", extract_structured_data)
    workflow.add_node("finalize", finalize_extraction)
    
    # Add edges
    workflow.add_edge(START, "determine_file_type")
    workflow.add_edge("determine_file_type", "extract_content")
    workflow.add_edge("extract_content", "detect_schema")
    workflow.add_edge("detect_schema", "extract_data")
    workflow.add_edge("extract_data", "finalize")
    workflow.add_edge("finalize", END)
    
    return workflow.compile()

# Initialize workflow
extraction_workflow = create_extraction_workflow()
if extraction_workflow is None:
    raise RuntimeError("Extraction workflow failed to initialize. Please check create_extraction_workflow and its dependencies.")

# API endpoints
@app.post("/extract", response_model=ExtractionResult)
async def extract_document_info(
    file: UploadFile = File(...),
    schema: Optional[str] = Form(None, description="JSON string of extraction schema (optional if auto_detect is True)"),
    auto_detect_schema: bool = Form(True, description="Whether to automatically detect the best schema (always enabled)")
):
    if extraction_workflow is None:
        raise HTTPException(status_code=500, detail="Extraction workflow is not initialized. Please contact the administrator.")
    """
    Extract structured information from uploaded documents
    
    Args:
        file: Uploaded file (PDF, Excel, or Word)
{{ ... }}
        schema: JSON string defining the fields to extract (optional if auto_detect_schema is True)
        auto_detect_schema: Whether to automatically detect the best schema
    
    Returns:
        Extracted structured data
    """
    start_time = datetime.now()
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Parse extraction schema if provided
        extraction_schema = None
        if schema:
            extraction_schema = json.loads(schema)
        elif not auto_detect_schema:
            raise HTTPException(status_code=400, detail="Either provide a schema or set auto_detect_schema=True")

        # --- ROUTE ALL FILES INCLUDING PDF THROUGH MAIN WORKFLOW ---
        # Create a temporary file for Excel files to use in enhanced extraction if needed
        excel_temp_path = None
        if file.filename and file.filename.lower().endswith(('.xlsx', '.xls')):
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as temp_file:
                temp_file.write(file_content)
                excel_temp_path = temp_file.name
                print(f"Created temporary Excel file: {excel_temp_path}")
        # Create initial state
        initial_state = DocumentProcessingState(
            file_name=file.filename,
            file_content=file_content,
            extraction_schema=extraction_schema,
            auto_detect_schema=auto_detect_schema,
            processing_stage="initialized"
        )
        # Run workflow (PDFs now use modular pipeline via extract_document_content)
        final_state = extraction_workflow.invoke(initial_state)
        processing_time = (datetime.now() - start_time).total_seconds()
        # Excel fix logic (as before)
        if excel_temp_path and file.filename and file.filename.lower().endswith(('.xlsx', '.xls')):
            # Get the extracted data
            if isinstance(final_state, dict):
                extracted_data = final_state.get("extracted_data", {})
            else:
                extracted_data = getattr(final_state, "extracted_data", {}) or {}
            extracted_data = final_state.get('extracted_data', {})
            extracted_sheets = [k for k in extracted_data.keys() if k != '_metadata']
            extraction_status = extracted_data.get('_metadata', {}).get('extraction_status', '')
            extraction_failed = (
                not extracted_sheets or 
                extraction_status == 'failure' or
                extraction_status == 'partial'
            )
            if extraction_failed:
                print(f"Regular extraction failed or incomplete, applying Excel fix...")
                try:
                    fixed_data = excel_fix.fix_excel_extraction(
                        file_path=excel_temp_path, 
                        original_filename=file.filename
                    )
                    if fixed_data:
                        print(f"Excel fix applied successfully, updating results")
                        final_state['extracted_data'] = fixed_data
                except Exception as e:
                    print(f"Excel fix failed: {str(e)}")
                    traceback.print_exc()
            # Clean up the temporary file
            try:
                if os.path.exists(excel_temp_path):
                    os.unlink(excel_temp_path)
                    print(f"Removed temporary file: {excel_temp_path}")
            except Exception as e:
                print(f"Error removing temp file: {str(e)}")
        else:
            # --- ORIGINAL WORKFLOW FOR NON-PDF FILES ---
            excel_temp_path = None
            # Create a temporary file for Excel files to use in enhanced extraction if needed
            if file.filename and file.filename.lower().endswith(('.xlsx', '.xls')):
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as temp_file:
                    temp_file.write(file_content)
                    excel_temp_path = temp_file.name
                    print(f"Created temporary Excel file: {excel_temp_path}")
            # Create initial state
            initial_state = DocumentProcessingState(
                file_name=file.filename,
                file_content=file_content,
                extraction_schema=extraction_schema,
                auto_detect_schema=auto_detect_schema,
                processing_stage="initialized"
            )
            # Run workflow
            final_state = extraction_workflow.invoke(initial_state)
            processing_time = (datetime.now() - start_time).total_seconds()
            # Excel fix logic (as before)
            if excel_temp_path and file.filename and file.filename.lower().endswith(('.xlsx', '.xls')):
                # Get the extracted data
                if isinstance(final_state, dict):
                    extracted_data = final_state.get("extracted_data", {})
                else:
                    extracted_data = getattr(final_state, "extracted_data", {}) or {}
                extracted_data = final_state.get('extracted_data', {})
                extracted_sheets = [k for k in extracted_data.keys() if k != '_metadata']
                extraction_status = extracted_data.get('_metadata', {}).get('extraction_status', '')
                extraction_failed = (
                    not extracted_sheets or 
                    extraction_status == 'failure' or
                    extraction_status == 'partial'
                )
                if extraction_failed:
                    print(f"Regular extraction failed or incomplete, applying Excel fix...")
                    try:
                        fixed_data = excel_fix.fix_excel_extraction(
                            file_path=excel_temp_path, 
                            original_filename=file.filename
                        )
                        if fixed_data:
                            print(f"Excel fix applied successfully, updating results")
                            final_state['extracted_data'] = fixed_data
                    except Exception as e:
                        print(f"Excel fix failed: {str(e)}")
                        traceback.print_exc()
                # Clean up the temporary file
                try:
                    if os.path.exists(excel_temp_path):
                        os.unlink(excel_temp_path)
                        print(f"Removed temporary file: {excel_temp_path}")
                except Exception as e:
                    print(f"Error removing temp file: {str(e)}")
        
        # Extract the data based on the type of result
        try:
            if isinstance(final_state, dict):
                extracted_data = final_state.get("extracted_data", {})
                file_type = final_state.get("file_type", "unknown")
                schema_detection = None
                if auto_detect_schema:
                    try:
                        # Patch: Defensive fallback for None values to avoid pydantic validation errors
                        detected_schema_type = final_state.get("detected_schema_type") or "unknown"
                        confidence_score = final_state.get("schema_confidence")
                        if confidence_score is None or not isinstance(confidence_score, (float, int)):
                            confidence_score = 0.0
                        suggested_schema = final_state.get("extraction_schema") or {}
                        reasoning = final_state.get("schema_reasoning") or "No reasoning provided"
                        schema_detection = SchemaDetectionResult(
                            detected_schema_type=detected_schema_type,
                            confidence_score=confidence_score,
                            suggested_schema=suggested_schema,
                            reasoning=reasoning
                        )
                    except Exception as e:
                        print(f"Warning: Could not create schema detection result: {str(e)}")
                        schema_detection = None
            else:
                # Handle object-based state
                extracted_data = getattr(final_state, "extracted_data", {}) or {}
                file_type = getattr(final_state, "file_type", "unknown")
                schema_detection = None
                if auto_detect_schema:
                    try:
                        # Patch: Defensive fallback for None values to avoid pydantic validation errors
                        detected_schema_type = getattr(final_state, "detected_schema_type", None) or "unknown"
                        confidence_score = getattr(final_state, "schema_confidence", None)
                        if confidence_score is None or not isinstance(confidence_score, (float, int)):
                            confidence_score = 0.0
                        suggested_schema = getattr(final_state, "extraction_schema", None) or {}
                        reasoning = getattr(final_state, "schema_reasoning", None) or "No reasoning provided"
                        schema_detection = SchemaDetectionResult(
                            detected_schema_type=detected_schema_type,
                            confidence_score=confidence_score,
                            suggested_schema=suggested_schema,
                            reasoning=reasoning
                        )
                    except Exception as e:
                        print(f"Warning: Could not create schema detection result: {str(e)}")
                        schema_detection = None
        except Exception as e:
            print(f"Warning: Error extracting data from final state: {str(e)}")
            extracted_data = {}
            file_type = "unknown"
            schema_detection = None
        
        # Create a directory for extraction results if it doesn't exist
        output_dir = "extraction_results"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate a filename based on the original filename and timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = os.path.splitext(file.filename)[0]
        output_filename = f"{original_name}_{timestamp}.json"
        output_path = os.path.join(output_dir, output_filename)
        
        # Determine overall extraction status based on individual sheets
        sheet_status = {}  # Initialize empty dict to prevent variable reference errors
        if hasattr(final_state, "sheet_status") and final_state.sheet_status:
            sheet_status = final_state.sheet_status
        elif isinstance(final_state, dict) and "sheet_status" in final_state:
            sheet_status = final_state.get("sheet_status", {})
        else:
            # Create status dict based on extracted data
            print("No sheet status found in state, creating based on extracted data")
            for sheet_name, sheet_data in extracted_data.items():
                if isinstance(sheet_data, dict) and len(sheet_data) > 1:  # More than just schema type
                    sheet_status[sheet_name] = "success"
                else:
                    sheet_status[sheet_name] = "failure"
        
        successful_sheets = [name for name, status in sheet_status.items() if status == "success"]
        
        # Add metadata about the extraction process
        extraction_metadata = {
            "_metadata": {
                "sheets": list(extracted_data.keys()),
                "successful_sheets": successful_sheets,
                "sheet_count": len(extracted_data),
                "successful_count": len(successful_sheets),
                "extraction_status": "success" if successful_sheets else "partial_failure"
            }
        }
        
        # If this is not a multi-sheet extraction, remove the metadata
        if len(extracted_data) <= 1 and not "_metadata" in extracted_data:
            extraction_result = extracted_data
        else:
            # For multi-sheet, include metadata and all extracted sheets
            extraction_result = {**extraction_metadata}
            for sheet_name, sheet_data in extracted_data.items():
                if sheet_name != "_metadata":
                    extraction_result[sheet_name] = sheet_data
        
        # For multi-sheet extraction, we may have multiple schema detections
        schema_info = None
        if schema_detection:
            try:
                schema_info = {
                    "detected_schema_type": "multiple_schemas",
                    "confidence_score": schema_detection.confidence_score,
                    "reasoning": "Multiple schemas detected for different sheets"
                }
            except Exception as e:
                print(f"Warning: Could not create schema detection info: {str(e)}")
                schema_info = {
                    "detected_schema_type": "multiple_schemas",
                    "confidence_score": 0.0,
                    "reasoning": "Multiple schemas detected for different sheets"
                }
        
        result_data = {
            "extracted_data": extraction_result,
            "schema_detection": schema_info,
            "processing_metadata": {
                "file_name": file.filename,
                "file_type": file_type,
                "processing_time": processing_time,
                "timestamp": datetime.now().isoformat(),
                "debug_files": os.path.join(os.path.abspath("debug_output"))
            }
        }
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result_data, f, indent=2, ensure_ascii=False)
        
        # Prepare response
        # Check if we have any successful extractions
        extraction_success = False
        
        # Debug output to help troubleshoot
        print(f"Checking extraction result: {type(extracted_data)}")
        if isinstance(extracted_data, dict):
            print(f"Extraction contains {len(extracted_data)} keys")
            
        # Handle empty extraction results
        if not extracted_data or (isinstance(extracted_data, dict) and len(extracted_data) == 0):
            print("WARNING: No extracted data available - completely empty")
            # Create a minimal empty structure instead of raising an error
            extracted_data = {
                "_metadata": {
                    "extraction_status": "error",
                    "error": "No data could be extracted"
                }
            }
        
        # Print some debug info about the extracted data
        if extracted_data and isinstance(extracted_data, dict):
            print(f"Extracted data has {len(extracted_data.keys())} keys (sheet names)")
            for sheet_name, sheet_data in extracted_data.items():
                print(f"Sheet: {sheet_name}")
        
        # Check if any sheet has actual data
        sheet_status = {}
        for sheet_name, sheet_data in extracted_data.items():
            if isinstance(sheet_data, dict) and len(sheet_data) > 0:  
                # Consider any data in the sheet as valid
                print(f"Sheet {sheet_name} has data: {list(sheet_data.keys())[:5]}")
                sheet_status[sheet_name] = "success"
            elif isinstance(sheet_data, list) and len(sheet_data) > 0:
                print(f"Sheet {sheet_name} has list data with {len(sheet_data)} items")
                sheet_status[sheet_name] = "success"
            else:
                sheet_status[sheet_name] = "error"
                # Ensure we have at least an empty dict for this sheet
                if not extracted_data.get(sheet_name):
                    extracted_data[sheet_name] = {}
        
        # Determine overall extraction status based on individual sheets
        successful_sheets = [name for name, status in sheet_status.items() if status == "success"]
        
        # Add metadata about the extraction process
        extraction_metadata = {
            "_metadata": {
                "sheets": list(extracted_data.keys()),
                "successful_sheets": successful_sheets,
                "sheet_count": len(extracted_data),
                "successful_count": len(successful_sheets),
                "extraction_status": "success" if successful_sheets else "partial_failure"
            }
        }
        
        # Properly combine metadata and sheet data for the final result
        extraction_result = {}
        
        # First add all sheet data
        for sheet_name, sheet_data in extracted_data.items():
            if sheet_name != "_metadata":
                extraction_result[sheet_name] = sheet_data
        
        # Then add the metadata
        if "_metadata" in extracted_data:
            # Use the existing _metadata if available
            extraction_result["_metadata"] = extracted_data["_metadata"]
        else:
            # Otherwise use our generated metadata
            extraction_result["_metadata"] = extraction_metadata["_metadata"]
            
        print(f"Combined {len(extracted_data)} sheets with metadata into final result")
        
        # Handle empty extraction gracefully
        if not extraction_result or len(extraction_result) <= 1 and "_metadata" in extraction_result:
            print("WARNING: Only metadata found in extraction result - adding placeholder")
            extraction_result["default"] = {
                "extraction_status": "no_data_extracted"
            }
        
        # Create a simplified schema detection without complex objects
        simplified_schema = None
        if schema_detection:
            try:
                # Just create a simple dictionary with basic info
                simplified_schema = {
                    "detected_schema_type": "multiple_schemas",
                    "confidence_score": 1.0,
                    "reasoning": "Detected individual schemas for each sheet"
                }
            except Exception as e:
                print(f"Warning: Could not create schema detection info: {str(e)}")
        
        # Debug the final extraction result before returning
        print(f"Final extraction result structure: {', '.join(extraction_result.keys() if isinstance(extraction_result, dict) else ['<not a dict>'])}")
        
        # Check if the extraction result is empty or only contains metadata
        if not extraction_result or (isinstance(extraction_result, dict) and 
                                     (len(extraction_result) == 0 or
                                      (len(extraction_result) <= 1 and "_metadata" in extraction_result))):
            print("WARNING: Empty or metadata-only extraction result detected, attempting to recover from debug files...")
            
            # Use a simpler approach to recovery that avoids scoping issues
            # We'll manually read the debug files and parse the JSON
            recovered_data = {}
            
            # Get debug files
            debug_dir = "debug_output"
            if os.path.exists(debug_dir):
                # List all text files in the debug directory
                debug_files = []
                for file in os.listdir(debug_dir):
                    if file.endswith(".txt"):
                        debug_files.append(os.path.join(debug_dir, file))
                
                print(f"Found {len(debug_files)} debug files to process")
                
                # Process each debug file
                for debug_file in debug_files:
                    try:
                        sheet_name = os.path.basename(debug_file).replace("_raw.txt", "").replace("_", " ")
                        print(f"Processing debug file for sheet: {sheet_name}")
                        
                        # Read the file content
                        with open(debug_file, "r", encoding="utf-8") as f:
                            content = f.read().strip()
                        
                        # Try to extract valid JSON
                        sheet_data = None
                        
                        # First try standard JSON parsing
                        try:
                            sheet_data = json.loads(content)
                            print(f"Successfully parsed debug file for sheet: {sheet_name}")
                        except json.JSONDecodeError:
                            # If that fails, try finding JSON with regex
                            try:
                                import re
                                json_match = re.search(r'\{[\s\S]*\}', content)
                                if json_match:
                                    sheet_data = json.loads(json_match.group(0))
                                    print(f"Successfully extracted JSON with regex for sheet: {sheet_name}")
                            except Exception:
                                print(f"Failed to extract valid JSON for sheet: {sheet_name}")
                        
                        # If we found data, add it to the recovered data
                        if sheet_data:
                            recovered_data[sheet_name] = sheet_data
                    except Exception as e:
                        print(f"Error processing debug file {debug_file}: {str(e)}")
                
                # If we recovered any data, create metadata and update the extraction result
                if recovered_data:
                    # Add metadata
                    recovered_data["_metadata"] = {
                        "sheets": [k for k in recovered_data.keys() if k != "_metadata"],
                        "extraction_status": "recovered_from_debug_files",
                        "recovery_timestamp": datetime.now().isoformat()
                    }
                    
                    # Update the extraction result
                    extraction_result = recovered_data
                    print(f"Successfully recovered data for {len(recovered_data) - 1} sheets")
                else:
                    print("No data could be recovered from debug files")
            else:
                print(f"Debug directory {debug_dir} not found")
                
            # Check if recovery was successful
            if extraction_result and isinstance(extraction_result, dict) and len(extraction_result) > 1:
                print(f"Recovery successful: {len(extraction_result)} keys in extraction result")
            else:
                print("Recovery failed - no valid extraction data available")
        
        # Ensure we're not losing sheet data when constructing the API response
        if isinstance(extraction_result, dict):
            sheet_count = 0
            for sheet_name in extraction_result.keys():
                if sheet_name != "_metadata":
                    sheet_count += 1
                    print(f"Final sheet data for {sheet_name}: {list(extraction_result[sheet_name].keys()) if isinstance(extraction_result[sheet_name], dict) else type(extraction_result[sheet_name])}")
            print(f"Final extraction contains {sheet_count} sheets with data")
        
        # For an API response, just return a simplified version
        return ExtractionResult(
            extracted_data=extraction_result,  # This should include all sheet data
            file_info={
                "filename": file.filename,
                "file_type": file_type,
                "size": len(file_content),
                "saved_to": output_path
            },
            processing_time=processing_time,
            status="success" if successful_sheets else "partial_success",
            schema_detection=simplified_schema  # Include schema detection info
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

@app.post("/detect-schema", response_model=SchemaDetectionResult)
async def detect_document_schema(file: UploadFile = File(...)):
    """
    Detect the most appropriate schema for a document without extracting data
    
    Args:
        file: Uploaded file (PDF, Excel, or Word)
    
    Returns:
        Schema detection result
    """
    try:
        # Read file content
        file_content = await file.read()
        
        # Create initial state for schema detection only
        initial_state = DocumentProcessingState(
            file_name=file.filename,
            file_content=file_content,
            auto_detect_schema=True,
            processing_stage="initialized"
        )
        
        # Run workflow up to schema detection
        state = determine_file_type(initial_state)
        state = extract_document_content(state)
        state = detect_schema(state)
        
        if state.error:
            raise HTTPException(status_code=500, detail=state.error)
        
        return SchemaDetectionResult(
            detected_schema_type=state.detected_schema_type,
            confidence_score=state.schema_confidence,
            suggested_schema=state.extraction_schema,
            reasoning=state.schema_reasoning
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Schema detection error: {str(e)}")

@app.get("/schemas")
async def get_available_schemas():
    """Get all available predefined schemas"""
    return {
        "available_schemas": {
            name: {
                "description": f"Schema for {name.replace('_', ' ').title()}",
                "fields": schema
            }
            for name, schema in PREDEFINED_SCHEMAS.items()
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.get("/supported-formats")
async def get_supported_formats():
    """Get supported file formats"""
    return {
        "supported_formats": {
            "pdf": [".pdf"],
            "excel": [".xlsx", ".xls"],
            "word": [".docx", ".doc"]
        }
    }

@app.post("/validate-schema")
async def validate_schema(schema: Dict[str, Any]):
    """Validate extraction schema, supporting nested structures"""
    def _validate(d):
        if not isinstance(d, dict):
            raise HTTPException(status_code=400, detail="Schema must be a dictionary at every level")
        if not d:
            raise HTTPException(status_code=400, detail="Schema (or nested schema) cannot be empty")
        for v in d.values():
            if isinstance(v, dict):
                _validate(v)
            elif isinstance(v, list):
                for item in v:
                    if isinstance(item, dict):
                        _validate(item)
    try:
        _validate(schema)
        return {"valid": True, "fields": list(schema.keys())}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Schema validation error: {str(e)}")

import errno

def _safe_float(val):
    try:
        return float(val) if val not in (None, "", "-") else 0.0
    except Exception:
        return 0.0

def postprocess_and_fill_missing_data(input_json_path: str, output_folder: str = "final_result"):
    """
    Post-process extraction results, fill missing/calculated fields, and save to output_folder.
    Follows user-provided calculation logic and schema.
    """
    import os
    import json
    os.makedirs(output_folder, exist_ok=True)
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    extracted_data = data.get('extracted_data', {})

    # --- Manufacturing Account ---
    mfg = extracted_data.get('Manufacturing Account', {})
    # cost_of_raw_material_consumed = opening_stock + purchases - returns_outwards - closing_stock
    rm = mfg.get('raw_material', {})
    opening_stock = _safe_float(rm.get('opening_stock'))
    purchases = _safe_float(rm.get('purchases'))
    returns_outwards = _safe_float(rm.get('returns_outwards'))
    closing_stock = _safe_float(rm.get('closing_stock'))
    cost_of_raw_material_consumed = opening_stock + purchases - returns_outwards - closing_stock
    rm['cost_of_raw_material_consumed'] = f"{cost_of_raw_material_consumed:.2f}"
    mfg['raw_material'] = rm

    # direct_labour.total = sum of bonus, casual_wages, epf, socso, eis, sub_contract_wages, wages_salaries
    dl = mfg.get('direct_labour', {})
    dl_total = sum(_safe_float(dl.get(key)) for key in [
        'bonus','casual_wages','epf','socso','eis','sub_contract_wages','wages_salaries'])
    dl['total'] = f"{dl_total:.2f}"
    mfg['direct_labour'] = dl

    # factory_overheads.total_overheads = depreciation + factory_expenses
    fo = mfg.get('factory_overheads', {})
    depreciation = _safe_float(fo.get('depreciation'))
    factory_expenses = _safe_float(fo.get('factory_expenses'))
    total_overheads = depreciation + factory_expenses
    fo['total_overheads'] = f"{total_overheads:.2f}"
    fo['total'] = f"{total_overheads:.2f}"
    mfg['factory_overheads'] = fo

    # total_cost = cost_of_raw_material_consumed + direct_labour.total + total_overheads
    total_cost = cost_of_raw_material_consumed + dl_total + total_overheads
    mfg['total_cost'] = f"{total_cost:.2f}"

    # production_cost = total_cost + opening_wip - closing_wip
    wip = mfg.get('work_in_progress', {})
    opening_wip = _safe_float(wip.get('opening'))
    closing_wip = _safe_float(wip.get('closing'))
    production_cost = total_cost + opening_wip - closing_wip
    mfg['production_cost'] = f"{production_cost:.2f}"
    extracted_data['Manufacturing Account'] = mfg

    # --- Trading P&L ---
    tpl = extracted_data.get('Trading P&L', {})
    # revenue = sales - returns_inwards
    sales = _safe_float(tpl.get('revenue'))
    returns_inwards = _safe_float(tpl.get('returns_inwards'))
    revenue = sales - returns_inwards
    tpl['revenue'] = f"{revenue:.2f}"
    # gross_profit = revenue - total cost of sales
    cos = tpl.get('cost_of_sales', {})
    cost_of_sales_total = _safe_float(cos.get('total'))
    gross_profit = revenue - cost_of_sales_total
    tpl['gross_profit'] = f"{gross_profit:.2f}"
    # other_income.total = sum of all other_income items (except total)
    oi = tpl.get('other_income', {})
    oi_total = sum(_safe_float(v) for k,v in oi.items() if k != 'total')
    oi['total'] = f"{oi_total:.2f}"
    tpl['other_income'] = oi
    # expenses.total = gross_profit + other_income.total - net_profit
    expenses = tpl.get('expenses', {})
    net_profit = _safe_float(tpl.get('net_profit'))
    expenses_total = gross_profit + oi_total - net_profit
    expenses['total'] = f"{expenses_total:.2f}"
    tpl['expenses'] = expenses
    # net_profit = gross_profit + other_income.total - expenses.total
    tpl['net_profit'] = f"{gross_profit + oi_total - expenses_total:.2f}"
    extracted_data['Trading P&L'] = tpl

    # --- Balance Sheet ---
    bs = extracted_data.get('Balance Sheet', {})
    # non_current_assets.total = property_plant_equipment + investments
    nca = bs.get('non_current_assets', {})
    ppe = _safe_float(nca.get('property_plant_equipment'))
    investments = _safe_float(nca.get('investments'))
    nca['total'] = f"{ppe + investments:.2f}"
    bs['non_current_assets'] = nca
    # current_assets.total = sum of all current asset items (except total)
    ca = bs.get('current_assets', {})
    ca_total = sum(_safe_float(v) for k,v in ca.items() if k != 'total')
    ca['total'] = f"{ca_total:.2f}"
    bs['current_assets'] = ca
    # current_liabilities.total = sum of all current liability items (except total)
    cl = bs.get('current_liabilities', {})
    cl_total = sum(_safe_float(v) for k,v in cl.items() if k != 'total')
    cl['total'] = f"{cl_total:.2f}"
    bs['current_liabilities'] = cl
    # non_current_liabilities.total = sum of all non-current liability items (except total)
    ncl = bs.get('non_current_liabilities', {})
    ncl_total = sum(_safe_float(v) for k,v in ncl.items() if k != 'total')
    ncl['total'] = f"{ncl_total:.2f}"
    bs['non_current_liabilities'] = ncl
    # equity.total = share_capital + retained_earnings
    eq = bs.get('equity', {})
    share_capital = _safe_float(eq.get('share_capital'))
    retained_earnings = _safe_float(eq.get('retained_earnings'))
    eq['total'] = f"{share_capital + retained_earnings:.2f}"
    bs['equity'] = eq
    # total_assets & total_liabilities_equity: take directly from document
    # (do not recalculate, as per user instruction)
    extracted_data['Balance Sheet'] = bs

    # Save to output folder
    base_name = os.path.basename(input_json_path)
    output_path = os.path.join(output_folder, base_name)
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Post-processed result saved to {output_path}")

# Example usage and testing
if __name__ == "__main__":
    import uvicorn
    
    print("Enhanced Document Information Extraction API with Schema Detection")
    print("Available predefined schemas:")
    for name, schema in PREDEFINED_SCHEMAS.items():
        print(f"  - {name}: {len(schema)} fields")
    
    print(f"\nAPI Features:")
    print("- Automatic schema detection using AI")
    print("- Manual schema specification")
    print("- Support for PDF, Excel, and Word documents")
    print("- Multiple predefined schemas for common document types")
    print("- Schema validation and management")
    
    print("\nStarting server...")
    uvicorn.run(app, host="0.0.0.0", port=8000)