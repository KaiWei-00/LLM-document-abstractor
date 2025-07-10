"""
Document Information Extraction Microservice using LangGraph
Supports PDF, Excel, and Word files with dynamic schema-based extraction
"""

import os
import json
import tempfile
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from datetime import datetime

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# FastAPI imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field

# LangGraph imports
from langgraph.graph import StateGraph, START, END
from langgraph.prebuilt import ToolNode
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langchain_core.tools import tool

# Document processing imports
import PyPDF2
import pandas as pd
import docx
import openpyxl
from io import BytesIO
import fitz  # PyMuPDF for better PDF handling

# Validate environment variables
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("OPENAI_API_KEY environment variable is required. Please set it in your .env file.")

# Initialize FastAPI app
app = FastAPI(title="Document Information Extraction API", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OpenAI LLM
try:
    # Check if the specified model is a completion model or a chat model
    model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
    if "instruct" in model_name:
        from langchain_openai import OpenAI
        llm = OpenAI(
            model=model_name,
            temperature=0,
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
    else:
        # Default to ChatOpenAI for other models
        llm = ChatOpenAI(
            model=model_name,
            temperature=0,
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
    print(f"Using OpenAI model: {model_name}")
except Exception as e:
    print(f"Error initializing OpenAI: {str(e)}")

# Pydantic models for API
class ExtractionSchema(BaseModel):
    """Define the schema for data extraction"""
    fields: Dict[str, str] = Field(description="Field names and their descriptions")
    
class DocumentType(BaseModel):
    """Supported document types"""
    pdf: bool = True
    excel: bool = True
    word: bool = True

class ExtractionResult(BaseModel):
    """Structured extraction result"""
    extracted_data: Dict[str, Any]
    file_info: Dict[str, Any]  # Now includes saved_to field with the JSON file path
    processing_time: float
    status: str

# State for LangGraph
class DocumentProcessingState(BaseModel):
    """State for document processing workflow"""
    file_content: Optional[Union[str, bytes]] = None
    file_type: Optional[str] = None
    file_name: Optional[str] = None
    extraction_schema: Optional[Dict[str, str]] = None
    extracted_data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    processing_stage: str = "initial"

# Document processing tools
@tool
def extract_pdf_content(file_path: str) -> str:
    """Extract text content from PDF files"""
    try:
        # Try with PyMuPDF first (better text extraction)
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e2:
            return f"Error extracting PDF: {str(e2)}"

@tool
def extract_excel_content(file_path: str) -> str:
    """Extract content from Excel files"""
    try:
        # Create a copy of the file to avoid locking issues
        import shutil
        import tempfile
        
        # Create a temporary file
        temp_dir = tempfile.gettempdir()
        temp_file = os.path.join(temp_dir, f"safe_copy_{os.path.basename(file_path)}")
        
        # Make a copy of the original file
        shutil.copy2(file_path, temp_file)
        
        # Process the copy
        content = ""
        
        # Use context manager to ensure file is closed
        with pd.ExcelFile(temp_file) as xl_file:
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(xl_file, sheet_name=sheet_name)
                content += f"\n--- Sheet: {sheet_name} ---\n"
                
                # Convert DataFrame to structured text
                content += df.to_string(index=False)
                content += "\n\n"
        
        # Clean up the temporary file
        try:
            os.unlink(temp_file)
        except:
            pass  # If we can't delete it now, it will be cleaned up later
            
        return content
    except Exception as e:
        return f"Error extracting Excel: {str(e)}"

@tool
def extract_word_content(file_path: str) -> str:
    """Extract content from Word documents"""
    try:
        doc = docx.Document(file_path)
        content = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
            
        # Extract tables
        for table in doc.tables:
            content += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                content += " | ".join(row_text) + "\n"
            content += "\n"
            
        return content
    except Exception as e:
        return f"Error extracting Word: {str(e)}"

# LangGraph workflow nodes
def determine_file_type(state: DocumentProcessingState) -> DocumentProcessingState:
    """Determine the file type based on extension"""
    if state.file_name:
        extension = Path(state.file_name).suffix.lower()
        if extension == '.pdf':
            state.file_type = 'pdf'
        elif extension in ['.xlsx', '.xls']:
            state.file_type = 'excel'
        elif extension in ['.docx', '.doc']:
            state.file_type = 'word'
        else:
            state.error = f"Unsupported file type: {extension}"
    
    state.processing_stage = "file_type_determined"
    return state

def extract_document_content(state: DocumentProcessingState) -> DocumentProcessingState:
    """Extract content based on file type"""
    if state.error:
        return state
    
    try:
        # Create temporary file with proper handling for bytes
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(state.file_name).suffix) as temp_file:
            if isinstance(state.file_content, bytes):
                temp_file.write(state.file_content)
            elif isinstance(state.file_content, str):
                temp_file.write(state.file_content.encode())
            else:
                state.error = f"Unsupported file_content type: {type(state.file_content)}"
                return state
            temp_path = temp_file.name
        
        # Extract based on file type
        if state.file_type == 'pdf':
            content = extract_pdf_content.invoke({"file_path": temp_path})
        elif state.file_type == 'excel':
            content = extract_excel_content.invoke({"file_path": temp_path})
        elif state.file_type == 'word':
            content = extract_word_content.invoke({"file_path": temp_path})
        else:
            state.error = f"Unsupported file type: {state.file_type}"
            return state
        
        state.file_content = content
        state.processing_stage = "content_extracted"
        
        # Clean up temporary file
        os.unlink(temp_path)
        
    except Exception as e:
        state.error = f"Error extracting content: {str(e)}"
    
    return state



def extract_structured_data(state: DocumentProcessingState) -> DocumentProcessingState:
    """Extract structured data using LLM based on schema"""
    if state.error:
        return state
    
    try:
        # Create extraction prompt
        schema_description = "\n".join([f"- {field}: {desc}" for field, desc in state.extraction_schema.items()])
        
        prompt = f"""
        Extract the following information from the document content below. 
        If a field is not found or cannot be determined, leave it as an empty string.
        
        Required fields:
        {schema_description}
        
        Document content:
        {state.file_content}
        
        Please return the extracted data in JSON format with the exact field names as keys.
        If any field is missing or cannot be found, use an empty string as the value.
        
        Example format:
        {{
            "company_name": "ABC Corp",
            "revenue": "1000000",
            "year": "2023",
            "missing_field": ""
        }}
        """
        
        # Use LLM to extract data
        messages = [
            SystemMessage(content="You are an expert document analyst. Extract structured data from documents according to the provided schema. Always return valid JSON."),
            HumanMessage(content=prompt)
        ]
        
        # Different handling for ChatOpenAI vs OpenAI models
        if isinstance(llm, ChatOpenAI):
            response = llm.invoke(messages)
            response_text = response.content
        else:
            # For completion-based models like gpt-3.5-turbo-instruct
            response = llm.invoke(prompt)
            response_text = response
        
        # Parse JSON response
        try:
            extracted_data = json.loads(response_text)
            state.extracted_data = extracted_data
            state.processing_stage = "data_extracted"
        except json.JSONDecodeError:
            # If JSON parsing fails, try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                extracted_data = json.loads(json_match.group())
                state.extracted_data = extracted_data
                state.processing_stage = "data_extracted"
            else:
                state.error = "Failed to parse extracted data as JSON"
        
    except Exception as e:
        state.error = f"Error extracting structured data: {str(e)}"
    
    return state

def finalize_extraction(state: DocumentProcessingState) -> DocumentProcessingState:
    """Finalize the extraction process"""
    if not state.error and state.extracted_data:
        # Ensure all schema fields are present
        for field in state.extraction_schema.keys():
            if field not in state.extracted_data:
                state.extracted_data[field] = ""
        
        state.processing_stage = "completed"
    
    return state

# Create LangGraph workflow
def create_extraction_workflow():
    """Create the document extraction workflow"""
    workflow = StateGraph(DocumentProcessingState)
    
    # Add nodes
    workflow.add_node("determine_file_type", determine_file_type)
    workflow.add_node("extract_content", extract_document_content)
    workflow.add_node("extract_data", extract_structured_data)
    workflow.add_node("finalize", finalize_extraction)
    
    # Add edges
    workflow.add_edge(START, "determine_file_type")
    workflow.add_edge("determine_file_type", "extract_content")
    workflow.add_edge("extract_content", "extract_data")
    workflow.add_edge("extract_data", "finalize")
    workflow.add_edge("finalize", END)
    
    return workflow.compile()

# Initialize workflow
extraction_workflow = create_extraction_workflow()

# API endpoints
@app.post("/extract", response_model=ExtractionResult)
async def extract_document_info(
    file: UploadFile = File(...),
    schema: str = Form(..., description="JSON string of extraction schema")
):
    """
    Extract structured information from uploaded documents
    
    Args:
        file: Uploaded file (PDF, Excel, or Word)
        schema: JSON string defining the fields to extract
    
    Returns:
        Extracted structured data
    """
    start_time = datetime.now()
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Parse extraction schema
        extraction_schema = json.loads(schema)
        
        # Create initial state
        initial_state = DocumentProcessingState(
            file_name=file.filename,
            file_content=file_content,
            extraction_schema=extraction_schema,
            processing_stage="initialized"
        )
        
        # Run workflow
        final_state = extraction_workflow.invoke(initial_state)
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Extract the data based on the type of result
        if isinstance(final_state, dict):
            if final_state.get("error"):
                raise HTTPException(status_code=500, detail=final_state["error"])
            extracted_data = final_state.get("extracted_data", {})
            file_type = final_state.get("file_type", "unknown")
        else:
            if final_state.error:
                raise HTTPException(status_code=500, detail=final_state.error)
            extracted_data = final_state.extracted_data
            file_type = final_state.file_type
        
        # Create a directory for extraction results if it doesn't exist
        output_dir = "extraction_results"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate a filename based on the original filename and timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = os.path.splitext(file.filename)[0]
        output_filename = f"{original_name}_{timestamp}.json"
        output_path = os.path.join(output_dir, output_filename)
        
        # Save the extracted data to a JSON file
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(extracted_data, f, indent=2, ensure_ascii=False)
        
        # Prepare response
        return ExtractionResult(
            extracted_data=extracted_data,
            file_info={
                "filename": file.filename,
                "file_type": file_type,
                "size": len(file_content),
                "saved_to": output_path
            },
            processing_time=processing_time,
            status="success"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.get("/supported-formats")
async def get_supported_formats():
    """Get supported file formats"""
    return {
        "supported_formats": {
            "pdf": [".pdf"],
            "excel": [".xlsx", ".xls"],
            "word": [".docx", ".doc"]
        }
    }

@app.post("/validate-schema")
async def validate_schema(schema: Dict[str, str]):
    """Validate extraction schema"""
    try:
        # Basic validation
        if not isinstance(schema, dict):
            raise HTTPException(status_code=400, detail="Schema must be a dictionary")
        
        if not schema:
            raise HTTPException(status_code=400, detail="Schema cannot be empty")
        
        return {"valid": True, "fields": list(schema.keys())}
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Schema validation error: {str(e)}")

# Example usage and testing
if __name__ == "__main__":
    import uvicorn
    
    # Example schema for Malaysian financial reports
    example_schema = {
        "company_name": "Name of the company",
        "financial_year": "Financial year of the report",
        "revenue": "Total revenue/sales",
        "net_profit": "Net profit/income",
        "total_assets": "Total assets",
        "total_liabilities": "Total liabilities",
        "cash_and_equivalents": "Cash and cash equivalents",
        "chairman_name": "Name of the chairman",
        "ceo_name": "Name of the CEO",
        "auditor_name": "Name of the auditing firm"
    }
    
    print("Example schema for Malaysian financial reports:")
    print(json.dumps(example_schema, indent=2))
    print("\nStarting server...")
    
    uvicorn.run(app, host="0.0.0.0", port=8000)